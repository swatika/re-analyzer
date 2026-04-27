"""
Real Estate Investment Analyzer — Austin TX
=============================================
Automated due diligence tool that pulls public data for any Austin address
and generates a comprehensive investment analysis report.

Usage:
    python analyze_deal.py "1309 Perez St" --zip 78721 --purchase-price 500000
    python analyze_deal.py "1309 Perez St" --zip 78721 --purchase-price 500000 --build-sf 6400 --exit-psf 450

Data Sources:
    - TCAD (Travis County Appraisal District) — appraisals, ownership, deeds
    - Austin Open Data API — construction permits
    - Redfin (via CSV API) — recent sold comps

Requirements:
    pip install requests python-docx openpyxl
"""

import argparse
import json
import re
import sys
import os
import subprocess
from datetime import datetime
from dataclasses import dataclass, field
from typing import Optional

import requests

# ── Data Classes ──────────────────────────────────────────────

@dataclass
class PropertyRecord:
    pid: int = 0
    address: str = ""
    legal: str = ""
    owner: str = ""
    market_value: float = 0
    land_value: float = 0
    improvement_value: float = 0
    acres: float = 0
    zoning: str = ""
    deeds: list = field(default_factory=list)

@dataclass
class Permit:
    address: str = ""
    description: str = ""
    sqft: float = 0
    issue_date: str = ""
    permit_class: str = ""
    work_class: str = ""
    builder: str = ""
    housing_units: int = 0
    floors: int = 0
    status: str = ""

@dataclass
class AnalysisResult:
    subject_properties: list = field(default_factory=list)
    nearby_properties: list = field(default_factory=list)
    street_permits: list = field(default_factory=list)
    zip_permits: list = field(default_factory=list)
    redfin_comps: list = field(default_factory=list)
    market_stats: dict = field(default_factory=dict)


# ── TCAD Module ───────────────────────────────────────────────

class TCADClient:
    """Travis County Appraisal District API client via TrueProdigy."""

    BASE_URL = "https://prod-container.trueprodigyapi.com"

    def __init__(self):
        self.token = None
        self._authenticate()

    def _authenticate(self):
        resp = requests.post(f"{self.BASE_URL}/trueprodigy/cadpublic/auth/token")
        resp.raise_for_status()
        self.token = resp.json()["user"]["token"]

    @property
    def headers(self):
        return {"Authorization": f"Bearer {self.token}"}

    def search(self, query: str) -> list:
        """Full-text property search."""
        resp = requests.get(
            f"{self.BASE_URL}/public/property/searchfulltext",
            params={"page": 1, "pageSize": 50},
            headers=self.headers,
            json={"searchText": query}
        )
        # The API uses POST-like GET with search text in body — fall back to Playwright
        if resp.status_code != 200:
            return self._search_via_playwright(query)
        data = resp.json()
        return data.get("results", [])

    def _search_via_playwright(self, query: str) -> list:
        """Use Playwright to search TCAD and capture API responses."""
        script = f"""
const {{ chromium }} = require('playwright');
(async () => {{
    const browser = await chromium.launch({{ headless: true }});
    const page = await browser.newPage();
    const results = [];
    page.on('response', async resp => {{
        try {{
            const url = resp.url();
            const ct = resp.headers()['content-type'] || '';
            if (ct.includes('json') && url.includes('searchfulltext')) {{
                const body = await resp.json();
                if (body.results) results.push(...body.results);
            }}
        }} catch(e) {{}}
    }});
    await page.goto('https://stage.travis.prodigycad.com/property-search', {{ waitUntil: 'networkidle', timeout: 30000 }});
    await page.waitForTimeout(2000);
    const input = await page.$('#searchInput');
    await input.fill('{query}');
    await page.keyboard.press('Enter');
    await page.waitForTimeout(5000);
    console.log(JSON.stringify(results));
    await browser.close();
}})();
"""
        result = subprocess.run(
            ["node", "-e", script],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0 and result.stdout.strip():
            return json.loads(result.stdout.strip())
        return []

    def get_deeds(self, pid: int) -> list:
        """Get deed history for a property via Playwright."""
        script = f"""
const {{ chromium }} = require('playwright');
(async () => {{
    const browser = await chromium.launch({{ headless: true }});
    const page = await browser.newPage();
    let deeds = [];
    page.on('response', async resp => {{
        try {{
            const url = resp.url();
            const ct = resp.headers()['content-type'] || '';
            if (ct.includes('json') && url.includes('/deeds')) {{
                const body = await resp.json();
                if (body.results) deeds = body.results;
            }}
        }} catch(e) {{}}
    }});
    await page.goto('https://stage.travis.prodigycad.com/property-detail/{pid}/2026', {{ waitUntil: 'networkidle', timeout: 20000 }});
    await page.waitForTimeout(3000);
    console.log(JSON.stringify(deeds));
    await browser.close();
}})();
"""
        result = subprocess.run(
            ["node", "-e", script],
            capture_output=True, text=True, timeout=45
        )
        if result.returncode == 0 and result.stdout.strip():
            try:
                return json.loads(result.stdout.strip())
            except json.JSONDecodeError:
                return []
        return []

    def search_with_deeds(self, query: str, zip_code: str = None) -> list[PropertyRecord]:
        """Search properties and get deed history for each."""
        raw = self._search_via_playwright(query)
        props = []
        seen = set()
        for r in raw:
            if r.get("pYear") != "2026":
                continue
            if zip_code and r.get("zip") != zip_code:
                continue
            pid = r.get("pid")
            if pid in seen:
                continue
            seen.add(pid)

            prop = PropertyRecord(
                pid=pid,
                address=r.get("streetPrimary", ""),
                legal=r.get("legalDescription", ""),
                owner=r.get("name", ""),
                market_value=r.get("marketValue", 0) or 0,
                land_value=r.get("landValue", 0) or 0,
                improvement_value=r.get("improvementValue", 0) or 0,
                acres=float(r.get("legalAcreage") or 0),
                zoning=r.get("zoning", ""),
            )
            props.append(prop)

        # Get deeds for each property
        print(f"  Fetching deed history for {len(props)} properties...")
        for prop in props:
            prop.deeds = self.get_deeds(prop.pid)

        return props


# ── Austin Permits Module ─────────────────────────────────────

class AustinPermits:
    """Austin Open Data API for construction permits."""

    BASE_URL = "https://data.austintexas.gov/resource/3syk-w9eu.json"

    def search_street(self, street_name: str, zip_code: str) -> list[Permit]:
        """Get new construction permits for a street."""
        where = f"permit_location like '%{street_name.upper()}%' AND permittype='BP' AND work_class='New' AND original_zip='{zip_code}'"
        params = {
            "$where": where,
            "$order": "issue_date DESC",
            "$limit": 100,
        }
        resp = requests.get(self.BASE_URL, params=params)
        if resp.status_code != 200:
            print(f"  Warning: Permit API returned {resp.status_code}")
            return []
        return self._parse_permits(resp.json())

    def search_zip(self, zip_code: str, limit: int = 200) -> list[Permit]:
        """Get all new residential construction permits in a zip code."""
        where = f"original_zip='{zip_code}' AND permittype='BP' AND work_class='New'"
        params = {
            "$where": where,
            "$order": "issue_date DESC",
            "$limit": limit,
        }
        resp = requests.get(self.BASE_URL, params=params)
        if resp.status_code != 200:
            return []
        permits = self._parse_permits(resp.json())
        # Filter residential only
        return [p for p in permits if 'Single Family' in p.permit_class
                or 'Two Family' in p.permit_class
                or 'Secondary' in p.permit_class]

    def search_address(self, address: str) -> list[Permit]:
        """Search permits for a specific address."""
        street_num = address.split()[0]
        street_name = ' '.join(address.split()[1:]).upper().replace(' ST', '').replace(' AVE', '').replace(' DR', '')
        where = f"permit_location like '%{street_num}%{street_name}%' AND permittype='BP'"
        params = {
            "$where": where,
            "$order": "issue_date DESC",
            "$limit": 50,
        }
        resp = requests.get(self.BASE_URL, params=params)
        if resp.status_code != 200:
            return []
        return self._parse_permits(resp.json())

    def _parse_permits(self, data: list) -> list[Permit]:
        permits = []
        for r in data:
            issue_date = ""
            if r.get("issue_date"):
                try:
                    dt = datetime.fromisoformat(r["issue_date"].replace("T", " ").split(".")[0])
                    issue_date = dt.strftime("%Y-%m-%d")
                except:
                    issue_date = str(r["issue_date"])[:10]

            permits.append(Permit(
                address=r.get("permit_location", ""),
                description=r.get("description", ""),
                sqft=float(r.get("total_new_add_sqft", 0) or 0),
                issue_date=issue_date,
                permit_class=r.get("permit_class", ""),
                work_class=r.get("work_class", ""),
                builder=r.get("contractor_company_name", "Unknown"),
                housing_units=int(r.get("housing_units", 0) or 0),
                floors=int(r.get("number_of_floors", 0) or 0),
                status=r.get("status_current", ""),
            ))
        return permits


# ── Redfin Comps Module ──────────────────────────────────────

class RedfinComps:
    """Get sold comps from Redfin using the CSV download API."""

    # Known zip-to-region_id mappings (Austin area)
    REGION_CACHE = {}

    def _get_region_id(self, zip_code: str) -> Optional[str]:
        """Look up Redfin region_id for a zip code."""
        if zip_code in self.REGION_CACHE:
            return self.REGION_CACHE[zip_code]
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36'
        }
        try:
            resp = requests.get(
                f'https://www.redfin.com/zipcode/{zip_code}',
                headers=headers, timeout=15
            )
            if resp.status_code == 200:
                # Extract regionId from page source  
                # Look for region_id=NNNNN pattern (not the zip itself)
                matches = re.findall(r'region_id=(\d+)', resp.text)
                for rid in matches:
                    if rid != zip_code and len(rid) >= 4:
                        self.REGION_CACHE[zip_code] = rid
                        return rid
                # Fallback: look for regionId in JSON
                match2 = re.search(r'"regionId"\s*:\s*(\d+)', resp.text)
                if match2 and match2.group(1) != zip_code:
                    rid = match2.group(1)
                    self.REGION_CACHE[zip_code] = rid
                    return rid
        except Exception:
            pass
        return None

    def get_sold_comps(self, zip_code: str) -> list[dict]:
        """Get recently sold new construction comps from Redfin CSV API."""
        region_id = self._get_region_id(zip_code)
        if not region_id:
            print(f"  ⚠ Could not find Redfin region_id for zip {zip_code}")
            return []

        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36'
        }
        url = (
            f'https://www.redfin.com/stingray/api/gis-csv?al=1&num_homes=200'
            f'&ord=redfin-recommended-asc&page_number=1'
            f'&region_id={region_id}&region_type=2'
            f'&sold_within_days=365&status=9&uipt=1&v=8&min_year_built=2020'
        )
        try:
            resp = requests.get(url, headers=headers, timeout=20)
            if resp.status_code != 200:
                print(f"  ⚠ Redfin CSV API returned status {resp.status_code}")
                return []

            lines = resp.text.strip().split('\n')
            # Skip disclaimer line if present
            header_idx = None
            for i, line in enumerate(lines):
                if line.startswith('SALE TYPE') or line.startswith('"SALE TYPE'):
                    header_idx = i
                    break
            if header_idx is None:
                return []

            import csv
            import io
            reader = csv.DictReader(io.StringIO('\n'.join(lines[header_idx:])))
            comps = []
            for row in reader:
                try:
                    price_str = (row.get('PRICE') or '0').replace(',', '').replace('$', '')
                    price = int(float(price_str)) if price_str else 0
                    sqft_str = (row.get('SQUARE FEET') or '0').replace(',', '')
                    sqft = int(float(sqft_str)) if sqft_str else 0
                    year_str = row.get('YEAR BUILT') or '0'
                    year = int(float(year_str)) if year_str else 0
                    psf = round(price / sqft) if sqft > 0 else 0
                    addr = row.get('ADDRESS') or ''
                    city = row.get('CITY') or ''
                    sold_date = row.get('SOLD DATE') or ''
                    beds = row.get('BEDS') or ''
                    baths = row.get('BATHS') or ''
                    if price > 0 and year >= 2020:
                        comps.append({
                            'address': f"{addr}, {city}",
                            'price': price,
                            'sqft': sqft,
                            'psf': psf,
                            'year_built': year,
                            'sold_date': sold_date,
                            'beds': beds,
                            'baths': baths,
                        })
                except (ValueError, ZeroDivisionError):
                    continue
            return comps
        except Exception as e:
            print(f"  ⚠ Redfin error: {e}")
            return []


# ── Analysis Engine ──────────────────────────────────────────

class InvestmentAnalyzer:
    """Main analysis engine that orchestrates all data sources."""

    def __init__(self):
        self.tcad = TCADClient()
        self.permits = AustinPermits()
        self.redfin = RedfinComps()

    def analyze(self, address: str, zip_code: str, street_name: str,
                purchase_price: float = 0, build_sf: float = 0,
                exit_psf: float = 0, build_cost_psf: float = 250) -> AnalysisResult:

        result = AnalysisResult()

        # 1. TCAD — Subject property
        print(f"\n[1/5] Searching TCAD for: {address}")
        result.subject_properties = self.tcad.search_with_deeds(address, zip_code)
        print(f"  Found {len(result.subject_properties)} subject property records")

        # 2. TCAD — Nearby properties on same street
        print(f"\n[2/5] Searching TCAD for all properties on {street_name} St...")
        result.nearby_properties = self.tcad.search_with_deeds(f"{street_name} ST", zip_code)
        print(f"  Found {len(result.nearby_properties)} properties on {street_name} St")

        # 3. Austin Permits — Street level
        print(f"\n[3/5] Pulling construction permits for {street_name} St, {zip_code}...")
        result.street_permits = self.permits.search_street(street_name, zip_code)
        print(f"  Found {len(result.street_permits)} new construction permits on {street_name} St")

        # 4. Austin Permits — Zip code level
        print(f"\n[4/5] Pulling all new residential permits in {zip_code}...")
        result.zip_permits = self.permits.search_zip(zip_code)
        print(f"  Found {len(result.zip_permits)} new residential permits in {zip_code}")

        # 5. Redfin sold comps
        print(f"\n[5/5] Scraping Redfin sold comps in {zip_code}...")
        result.redfin_comps = self.redfin.get_sold_comps(zip_code)
        print(f"  Found {len(result.redfin_comps)} sold comps")

        # Compute market stats
        if result.redfin_comps:
            psf_values = [c["psf"] for c in result.redfin_comps if c.get("psf", 0) > 0]
            if psf_values:
                psf_values.sort()
                mid = len(psf_values) // 2
                result.market_stats = {
                    "median_psf": psf_values[mid],
                    "avg_psf": round(sum(psf_values) / len(psf_values)),
                    "min_psf": min(psf_values),
                    "max_psf": max(psf_values),
                    "count": len(psf_values),
                }

        return result


# ── Report Generator ─────────────────────────────────────────

def generate_report(result: AnalysisResult, address: str, zip_code: str,
                    street_name: str, purchase_price: float,
                    build_sf: float, exit_psf: float,
                    build_cost_psf: float, output_path: str):
    """Generate a formatted Word document report."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    def set_cell_shading(cell, color):
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    def add_table(headers, rows, header_color='1F4E79'):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_shading(cell, header_color)
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = str(val)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)
                if r_idx % 2 == 1:
                    set_cell_shading(cell, 'F2F2F2')
        return table

    def fmt(val):
        if val is None:
            return "N/A"
        return f"${val:,.0f}"

    # ── Title Page ──
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f'{address}, {zip_code}')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Investment Due Diligence Report')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(89, 89, 89)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(datetime.now().strftime('%B %Y'))
    run.font.size = Pt(14)
    doc.add_page_break()

    # ── 1. Subject Property ──
    doc.add_heading('1. Subject Property — TCAD Data', level=1)
    if result.subject_properties:
        rows = []
        total_market = 0
        for p in result.subject_properties:
            total_market += p.market_value
            is_dev = any(x in p.owner.upper() for x in ['LLC', 'LP', 'INC', 'TRUST', 'CORP'])
            rows.append([p.address, str(p.pid), p.owner,
                         'Developer' if is_dev else 'Individual',
                         p.zoning, fmt(p.land_value), fmt(p.improvement_value),
                         fmt(p.market_value), f"{p.acres:.4f}"])
        add_table(
            ['Address', 'Prop ID', 'Owner', 'Type', 'Zoning', 'Land', 'Improvements', 'TCAD Market', 'Acres'],
            rows
        )
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(f'TCAD Total Market Value: {fmt(total_market)}')
        run.bold = True
        if purchase_price > 0:
            diff = purchase_price - total_market
            pct = (diff / total_market) * 100 if total_market > 0 else 0
            p.add_run(f' | Purchase Price: {fmt(purchase_price)} ({pct:+.1f}% vs TCAD)')

        # Deed history
        doc.add_heading('Deed History', level=2)
        for prop in result.subject_properties:
            if prop.deeds:
                doc.add_paragraph(f'{prop.address} (PID: {prop.pid}):')
                deed_rows = []
                for d in prop.deeds:
                    deed_rows.append([
                        d.get("deedDt", "")[:10],
                        d.get("deedDescription", d.get("deedType", "")),
                        d.get("seller", "") or "",
                        d.get("buyer", "") or "",
                    ])
                add_table(['Date', 'Type', 'From', 'To'], deed_rows)
                doc.add_paragraph()
    else:
        doc.add_paragraph('No TCAD records found for this address.')

    doc.add_page_break()

    # ── 2. Neighborhood Properties ──
    doc.add_heading(f'2. {street_name} St — All Properties & Deed Records', level=1)
    if result.nearby_properties:
        # Separate developers vs individuals
        developers = [p for p in result.nearby_properties
                      if any(x in p.owner.upper() for x in ['LLC', 'LP', 'INC', 'FUND'])]
        individuals = [p for p in result.nearby_properties
                       if p not in developers]

        doc.add_heading('Developer-Owned Properties', level=2)
        if developers:
            rows = []
            for p in developers:
                last_deed = ""
                if p.deeds:
                    wd = [d for d in p.deeds if d.get("deedType") in ("WD", "SW")]
                    if wd:
                        last_deed = f"{wd[0].get('deedDt', '')[:10]}"
                rows.append([p.address, p.owner, fmt(p.market_value),
                             fmt(p.land_value), fmt(p.improvement_value), last_deed])
            add_table(['Address', 'Developer', 'TCAD Market', 'Land', 'Improvements', 'Last Transfer'],
                      rows, header_color='8B0000')
        else:
            doc.add_paragraph('None found.')

        doc.add_heading('Individual-Owned Properties', level=2)
        rows = []
        for p in sorted(individuals, key=lambda x: x.address):
            rows.append([p.address, p.owner, fmt(p.market_value),
                         fmt(p.land_value), fmt(p.improvement_value)])
        if rows:
            add_table(['Address', 'Owner', 'TCAD Market', 'Land', 'Improvements'], rows)

    doc.add_page_break()

    # ── 3. Construction Permits ──
    doc.add_heading(f'3. Construction Permits — {street_name} St', level=1)
    if result.street_permits:
        # Group by status
        active = [p for p in result.street_permits if p.status == 'Active']
        final = [p for p in result.street_permits if p.status == 'Final']

        if active:
            doc.add_heading('Currently Under Construction (Active Permits)', level=2)
            rows = [[p.address, f"{p.sqft:,.0f} sf", p.issue_date, p.builder,
                     p.permit_class, p.description[:80]]
                    for p in active]
            add_table(['Address', 'Size', 'Permit Date', 'Builder', 'Type', 'Description'],
                      rows, header_color='D4A017')

        if final:
            doc.add_heading('Completed Projects (Final Permits)', level=2)
            rows = [[p.address, f"{p.sqft:,.0f} sf", p.issue_date, p.builder,
                     p.permit_class, p.description[:80]]
                    for p in final]
            add_table(['Address', 'Size', 'Permit Date', 'Builder', 'Type', 'Description'], rows)

    doc.add_heading(f'Zip Code {zip_code} — New Residential Permit Summary', level=2)
    doc.add_paragraph(f'Total new residential permits found: {len(result.zip_permits)}')
    if result.zip_permits:
        # By year
        by_year = {}
        for p in result.zip_permits:
            yr = p.issue_date[:4] if p.issue_date else 'Unknown'
            by_year.setdefault(yr, []).append(p)
        rows = []
        for yr in sorted(by_year.keys(), reverse=True):
            permits = by_year[yr]
            total_sf = sum(p.sqft for p in permits)
            rows.append([yr, str(len(permits)), f"{total_sf:,.0f} sf",
                         f"{total_sf / len(permits):,.0f} sf" if permits else "0"])
        add_table(['Year', 'Permits', 'Total SF', 'Avg SF/Permit'], rows)

    doc.add_page_break()

    # ── 4. Redfin Sold Comps ──
    doc.add_heading(f'4. Redfin Sold Comps — {zip_code} (New Construction)', level=1)
    if result.redfin_comps:
        stats = result.market_stats
        doc.add_heading('Market Statistics', level=2)
        add_table(
            ['Metric', 'Value'],
            [
                ['Median $/sf', f"${stats.get('median_psf', 0)}/sf"],
                ['Average $/sf', f"${stats.get('avg_psf', 0)}/sf"],
                ['Min $/sf', f"${stats.get('min_psf', 0)}/sf"],
                ['Max $/sf', f"${stats.get('max_psf', 0)}/sf"],
                ['Count', str(stats.get('count', 0))],
            ]
        )

        doc.add_heading('Individual Comps', level=2)
        comps_sorted = sorted(result.redfin_comps, key=lambda x: x.get("psf", 0), reverse=True)
        rows = [[c["address"], fmt(c.get("price", 0)),
                 f"{c.get('sqft', 0):,} sf", f"${c.get('psf', 0)}/sf"]
                for c in comps_sorted if c.get("psf", 0) > 0]
        if rows:
            add_table(['Address', 'Sold Price', 'Size', '$/sf'], rows)
    else:
        doc.add_paragraph('No Redfin comps retrieved. Run manually or check Playwright setup.')

    doc.add_page_break()

    # ── 5. Investment Analysis ──
    doc.add_heading('5. Investment Analysis & Recommendation', level=1)

    if purchase_price > 0 and build_sf > 0:
        total_build_cost = build_cost_psf * build_sf
        median_psf = result.market_stats.get("median_psf", 0)

        doc.add_heading('Deal Parameters', level=2)
        add_table(
            ['Parameter', 'Value'],
            [
                ['Purchase Price', fmt(purchase_price)],
                ['Build Size', f"{build_sf:,.0f} sf"],
                ['Build Cost @ ${:,.0f}/sf'.format(build_cost_psf), fmt(total_build_cost)],
                ['Exit Price Assumption', f"${exit_psf}/sf" if exit_psf else "N/A"],
                ['Market Median (Redfin)', f"${median_psf}/sf" if median_psf else "N/A"],
            ]
        )

        if exit_psf and median_psf:
            gap = ((exit_psf - median_psf) / median_psf) * 100
            doc.add_paragraph()
            p = doc.add_paragraph()
            if gap > 15:
                run = p.add_run(f'⚠ EXIT PRICE WARNING: ')
                run.bold = True
                run.font.color.rgb = RGBColor(192, 0, 0)
                p.add_run(f'${exit_psf}/sf is {gap:.0f}% above market median (${median_psf}/sf). High risk.')
            elif gap > 5:
                run = p.add_run(f'⚠ EXIT PRICE CAUTION: ')
                run.bold = True
                run.font.color.rgb = RGBColor(200, 150, 0)
                p.add_run(f'${exit_psf}/sf is {gap:.0f}% above market median (${median_psf}/sf). Moderate risk.')
            else:
                run = p.add_run(f'✅ EXIT PRICE REASONABLE: ')
                run.bold = True
                run.font.color.rgb = RGBColor(0, 128, 0)
                p.add_run(f'${exit_psf}/sf is within {gap:.0f}% of market median (${median_psf}/sf).')

        # Scenario table
        doc.add_heading('Exit Scenarios', level=2)
        scenarios = []
        for psf in [median_psf, median_psf + 25, median_psf + 50, exit_psf] if exit_psf else [median_psf, median_psf + 25, median_psf + 50]:
            if psf > 0:
                revenue = psf * build_sf
                total_cost = purchase_price + total_build_cost + (total_build_cost * 0.10)  # 10% soft/contingency
                profit = revenue - total_cost
                margin = (profit / total_cost) * 100
                scenarios.append([
                    f"${psf}/sf",
                    fmt(revenue),
                    fmt(total_cost),
                    fmt(profit),
                    f"{margin:.1f}%",
                    "✅" if margin > 10 else ("⚠" if margin > 0 else "❌")
                ])
        add_table(['Exit $/sf', 'Revenue', 'Total Cost', 'Profit', 'Margin', ''], scenarios)

    # Risk factors
    doc.add_heading('Risk Assessment', level=2)

    risks = []
    # Check for foreclosures
    for p in result.nearby_properties:
        for d in p.deeds:
            if d.get("deedType") == "ST":  # Substitute Trustee = foreclosure
                risks.append(f"FORECLOSURE on street: {p.address} — {d.get('deedDt', '')[:10]}")

    # Check for developer-held unsold inventory
    unsold = [p for p in result.nearby_properties
              if any(x in p.owner.upper() for x in ['LLC', 'LP', 'FUND'])
              and p.improvement_value > 100000]
    if unsold:
        total_unsold = sum(p.market_value for p in unsold)
        risks.append(f"${total_unsold:,.0f} in developer-held unsold inventory ({len(unsold)} properties)")

    # Active construction = competing supply
    active_permits = [p for p in result.street_permits if p.status == 'Active']
    if active_permits:
        risks.append(f"{len(active_permits)} competing units currently under construction on {street_name} St")

    # TCAD vs purchase price
    if result.subject_properties and purchase_price > 0:
        tcad_total = sum(p.market_value for p in result.subject_properties)
        if purchase_price > tcad_total * 1.05:
            risks.append(f"Purchase price ({fmt(purchase_price)}) is above TCAD appraisal ({fmt(tcad_total)}) — no equity cushion")

    for risk in risks:
        p = doc.add_paragraph()
        run = p.add_run(f'🔴 {risk}')
        run.font.color.rgb = RGBColor(192, 0, 0)

    if not risks:
        doc.add_paragraph('No major red flags identified from public data.')

    # ── Recommendation ──
    doc.add_heading('Recommendation', level=2)
    p = doc.add_paragraph()
    red_flag_count = len(risks)
    if red_flag_count >= 3:
        run = p.add_run('❌ NOT RECOMMENDED as currently structured. ')
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(192, 0, 0)
        p.add_run(f'{red_flag_count} risk factors identified. Re-evaluate pricing assumptions and market conditions.')
    elif red_flag_count >= 1:
        run = p.add_run('⚠ PROCEED WITH CAUTION. ')
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(200, 150, 0)
        p.add_run(f'{red_flag_count} risk factor(s) identified. Validate assumptions before committing.')
    else:
        run = p.add_run('✅ APPEARS VIABLE based on public data. ')
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 128, 0)
        p.add_run('Proceed with standard due diligence.')

    # ── Appendix ──
    doc.add_page_break()
    doc.add_heading('Appendix: Data Sources', level=1)
    sources = [
        ('TCAD', 'Travis County Appraisal District — ProdigyCAD public portal (stage.travis.prodigycad.com)'),
        ('Austin Open Data', 'Issued Construction Permits (data.austintexas.gov/resource/3syk-w9eu)'),
        ('Redfin', 'Recently sold new construction comps — scraped via Playwright'),
    ]
    for name, desc in sources:
        p = doc.add_paragraph()
        run = p.add_run(f'{name}: ')
        run.bold = True
        p.add_run(desc)

    p = doc.add_paragraph()
    run = p.add_run('\nDisclaimer: ')
    run.bold = True
    p.add_run('Texas is a non-disclosure state. Actual sale prices are not in public deed records. '
              'TCAD appraisals and MLS/listing data are proxies. This report is for informational purposes only.')

    doc.save(output_path)
    print(f"\n✅ Report saved: {output_path}")


# ── CLI Entry Point ──────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Austin TX Real Estate Investment Analyzer",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python analyze_deal.py "1309 Perez St" --zip 78721
  python analyze_deal.py "1309 Perez St" --zip 78721 --purchase-price 500000
  python analyze_deal.py "1309 Perez St" --zip 78721 --purchase-price 500000 --build-sf 6400 --exit-psf 450
        """
    )
    parser.add_argument("address", help="Property address (e.g., '1309 Perez St')")
    parser.add_argument("--zip", required=True, help="ZIP code (e.g., 78721)")
    parser.add_argument("--purchase-price", type=float, default=0, help="Purchase price in dollars")
    parser.add_argument("--build-sf", type=float, default=0, help="Total planned build square footage")
    parser.add_argument("--exit-psf", type=float, default=0, help="Expected exit price per square foot")
    parser.add_argument("--build-cost-psf", type=float, default=250, help="Build cost per sf (default: $250)")
    parser.add_argument("--output", default=None, help="Output path for Word doc")

    args = parser.parse_args()

    # Extract street name from address
    parts = args.address.upper().replace(",", "").split()
    # Remove street number and suffix
    suffixes = {"ST", "AVE", "DR", "LN", "BLVD", "CT", "WAY", "RD", "CIR", "PL"}
    street_parts = [p for p in parts[1:] if p not in suffixes]
    street_name = " ".join(street_parts) if street_parts else parts[1] if len(parts) > 1 else parts[0]

    output = args.output or os.path.join(
        os.path.expanduser("~"), "Downloads",
        f"{street_name.title()}_St_Investment_Report.docx"
    )

    print("=" * 60)
    print(f"  Real Estate Investment Analyzer — Austin TX")
    print(f"  Address:  {args.address}")
    print(f"  ZIP:      {args.zip}")
    print(f"  Street:   {street_name} St")
    if args.purchase_price:
        print(f"  Purchase: ${args.purchase_price:,.0f}")
    if args.build_sf:
        print(f"  Build:    {args.build_sf:,.0f} sf @ ${args.build_cost_psf:,.0f}/sf")
    if args.exit_psf:
        print(f"  Exit:     ${args.exit_psf:,.0f}/sf")
    print("=" * 60)

    analyzer = InvestmentAnalyzer()
    result = analyzer.analyze(
        address=args.address,
        zip_code=args.zip,
        street_name=street_name,
        purchase_price=args.purchase_price,
        build_sf=args.build_sf,
        exit_psf=args.exit_psf,
        build_cost_psf=args.build_cost_psf,
    )

    generate_report(
        result=result,
        address=args.address,
        zip_code=args.zip,
        street_name=street_name,
        purchase_price=args.purchase_price,
        build_sf=args.build_sf,
        exit_psf=args.exit_psf,
        build_cost_psf=args.build_cost_psf,
        output_path=output,
    )

    # Print summary to console
    print("\n" + "=" * 60)
    print("  QUICK SUMMARY")
    print("=" * 60)

    if result.subject_properties:
        total = sum(p.market_value for p in result.subject_properties)
        print(f"  TCAD Total Appraisal:  ${total:,.0f}")

    if result.market_stats:
        print(f"  Redfin Median $/sf:    ${result.market_stats.get('median_psf', 0)}/sf")
        print(f"  Redfin Avg $/sf:       ${result.market_stats.get('avg_psf', 0)}/sf")

    print(f"  Street Permits:        {len(result.street_permits)} new construction")
    active = len([p for p in result.street_permits if p.status == 'Active'])
    print(f"  Active (under const.): {active}")
    print(f"  Zip Permits:           {len(result.zip_permits)} residential")
    print(f"  Redfin Comps:          {len(result.redfin_comps)} sold")
    print(f"\n  Report: {output}")
    print("=" * 60)


if __name__ == "__main__":
    main()
