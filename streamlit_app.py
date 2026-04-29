"""
Austin Deal Analyzer PRO — Streamlit Web App
=============================================
Combined due diligence + financial modeling + BUY/DON'T BUY recommendation.
Pulls real market data, runs scenario analysis, and gives a clear verdict.
"""

import streamlit as st
import requests
import re
import io
import csv
import time
import numpy as np
from datetime import datetime, timedelta
from dataclasses import dataclass, field
from typing import Optional
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, numbers, Border, Side

# ── Gemini AI (optional, uses REST API — no extra package needed) ──
GEMINI_AVAILABLE = True  # Always available since we use REST API


def _get_gemini_key():
    try:
        return st.secrets.get("GEMINI_API_KEY", "")
    except Exception:
        return ""


def generate_ai_summary(deal_data: dict) -> str:
    """Generate a plain-English AI deal analysis using Google Gemini REST API (no SDK needed)."""
    api_key = _get_gemini_key()
    if not api_key:
        return ""

    prompt = f"""You are a real estate investment analyst. Analyze this deal and give a clear, 
plain-English summary that a beginner investor (first-time flipper) can understand. Be specific with numbers.

You MUST answer ALL of these questions clearly with specific numbers:

## 1. SHOULD I BUY THIS DEAL?
Give a clear YES or NO with reasoning. Consider the profit margin, exit price vs comps, and risks.

## 2. FLIP OR HOLD — WHICH STRATEGY IS BETTER?
- **Flip (sell immediately after build):** What's the projected profit? Is the exit price realistic?
- **Hold & Rent:** What monthly rent should I charge? How long should I hold before selling? What's the cash flow?
- **Recommend one strategy** and explain why.

## 3. WHAT PRICE SHOULD I TARGET FOR RENTING?
- Based on Census median rents and the property size, suggest a specific monthly rent per unit.
- Census median rent for ZIP {deal_data.get('zip_code', 'N/A')}: {deal_data.get('census_median_rent', 'N/A')}/mo
- Census rent by bedrooms: {deal_data.get('census_rent_by_br', 'N/A')}
- How long to hold if renting? Give a specific timeframe (e.g., 2 years, 5 years).

## 4. WHAT SHOULD MY MAXIMUM PURCHASE PRICE BE?
- Based on the comps and desired profit margin (20%+), what is the max I should pay for this land/property?
- Show the math: Max Purchase = Exit Revenue - Build Costs - Desired Profit

## 5. WHAT EXIT PRICE PER SQUARE FOOT IS REALISTIC?
- Based on nearby sold comps, what $/sf should I realistically expect?
- Is the current exit assumption of {deal_data.get('exit_psf', 0)}/sf achievable?

CRITICAL CONTEXT: This is a fix-and-flip deal. Negative monthly rental cash flow is normal for flips — rentals are a backup plan, not the primary strategy. Focus on flip profit first.

IMPORTANT FORMATTING RULES: Do NOT use dollar signs ($) for currency. Instead write amounts like "450,000" or "575/sf". Do NOT use LaTeX or math notation. Use plain text only.

DEAL DATA:
- Address: {deal_data.get('address', 'N/A')}, ZIP: {deal_data.get('zip_code', 'N/A')}
- Purchase Price: ${deal_data.get('purchase_price', 0):,.0f}
- Total Build Size: {deal_data.get('total_sf', 0):,} sq ft ({deal_data.get('num_units', 1)} units, {deal_data.get('per_unit_sf', 0):,.0f} sf/unit)
- Exit Price/SF: ${deal_data.get('exit_psf', 0)}/sf
- Total All-In Cost: ${deal_data.get('total_cost', 0):,.0f}
- Expected Revenue: ${deal_data.get('revenue', 0):,.0f}
- Projected Profit: ${deal_data.get('profit', 0):,.0f}
- Profit Margin: {deal_data.get('margin_pct', 0):.1f}%
- Break-Even PSF: ${deal_data.get('breakeven_psf', 0):.0f}/sf
- Market Median PSF: ${deal_data.get('median_psf', 0)}/sf (from {deal_data.get('comp_count', 0)} comps)
- Risk Score: {deal_data.get('risk_score', 0)}/100
- Verdict: {deal_data.get('verdict', 'N/A')}
- Listing Status: {deal_data.get('listing_status', 'Unknown')}
- Zoning: {deal_data.get('zoning', 'N/A')}
- Monthly Rent: ${deal_data.get('monthly_rent', 0):,.0f}
- Rental NOI/Year: ${deal_data.get('rental_noi', 0):,.0f}

COMPARABLE SALES (nearby sold properties):
{deal_data.get('top_comps', 'No comp data')}

CONSTRUCTION ACTIVITY:
- Active permits on street: {deal_data.get('active_permits', 0)}
- Completed projects: {deal_data.get('completed_permits', 0)}
- Notable builders: {deal_data.get('permit_builders', 'N/A')}

HOLD/RENTAL ANALYSIS ({deal_data.get('hold_months', 0)} month hold):
- Monthly NOI: ${deal_data.get('monthly_noi', 0):,.0f}
- Monthly Cash Flow After Debt: ${deal_data.get('monthly_cf_after_debt', 0):,.0f}
- Equity Multiple: {deal_data.get('equity_multiple', 0):.2f}x
- Annualized Return: {deal_data.get('annualized_return', 0):.1f}%
- Cash-on-Cash Yield: {deal_data.get('cash_yield', 0):.1f}%

FINANCING:
- Loan Amount: ${deal_data.get('loan_amount', 0):,.0f}
- Construction Interest: ${deal_data.get('construction_interest', 0):,.0f}
- Build Cost: ${deal_data.get('build_cost_psf', 0)}/sf
- Build Timeline: {deal_data.get('build_months', 0)} months

Keep your response under 800 words. Use headers (##) for each section. Use bullet points for clarity.
End with a FINAL VERDICT: BUY or DON'T BUY, and your recommended strategy (Flip or Hold).
The app's automated verdict is: {deal_data.get('verdict', 'N/A')}. Your recommendation should be consistent with this."""

    models = ["gemini-2.5-flash-lite", "gemini-2.5-flash", "gemini-2.0-flash"]
    last_error = ""
    try:
        for model_name in models:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={api_key}"
            try:
                resp = requests.post(url, json={
                    "contents": [{"parts": [{"text": prompt}]}]
                }, timeout=30)
            except Exception:
                continue
            if resp.status_code == 200:
                data = resp.json()
                return data["candidates"][0]["content"]["parts"][0]["text"]
            elif resp.status_code in (429, 404, 503):
                last_error = f"HTTP {resp.status_code}: {resp.text[:200]}"
                time.sleep(1)
                continue  # try next model
            else:
                last_error = f"HTTP {resp.status_code}: {resp.text[:200]}"
                continue  # try next model too
        # All models failed
        if "429" in last_error:
            return ("⚠️ **Gemini API quota exceeded.** The free tier has a daily limit.\n\n"
                    "**Options:**\n"
                    "- Wait until tomorrow (quota resets daily)\n"
                    "- Create a new API key at [Google AI Studio](https://aistudio.google.com/apikey) "
                    "and update it in Streamlit Cloud → Settings → Secrets")
        return f"⚠️ AI analysis error — all models failed. Last error: {last_error}"
    except Exception as e:
        return f"⚠️ AI analysis unavailable: {str(e)}"


# ── Page Config ──
st.set_page_config(
    page_title="Austin Deal Analyzer PRO",
    page_icon="🏡",
    layout="wide",
)

# ── Data Classes ──

@dataclass
class Permit:
    address: str = ""
    description: str = ""
    sqft: float = 0
    issue_date: str = ""
    permit_class: str = ""
    work_class: str = ""
    builder: str = ""
    contractor_name: str = ""
    applicant_name: str = ""
    applicant_org: str = ""
    housing_units: int = 0
    floors: int = 0
    status: str = ""
    permit_number: str = ""
    permit_type: str = ""
    permit_type_desc: str = ""

@dataclass
class AnalysisResult:
    street_permits: list = field(default_factory=list)
    street_all_permits: list = field(default_factory=list)
    zip_permits: list = field(default_factory=list)
    redfin_comps: list = field(default_factory=list)
    neighborhood_comps: list = field(default_factory=list)
    active_comps: list = field(default_factory=list)
    rental_comps: list = field(default_factory=list)
    rental_stats: dict = field(default_factory=dict)
    market_stats: dict = field(default_factory=dict)
    neighborhood_stats: dict = field(default_factory=dict)
    active_stats: dict = field(default_factory=dict)
    sources_status: dict = field(default_factory=dict)
    listing_status: dict = field(default_factory=dict)


# ── Austin Permits Module ──

class AustinPermits:
    BASE_URL = "https://data.austintexas.gov/resource/3syk-w9eu.json"

    def search_street(self, street_name: str, zip_code: str) -> list[Permit]:
        two_years_ago = (datetime.now() - timedelta(days=730)).strftime('%Y-%m-%dT00:00:00')
        where = f"permit_location like '%{street_name.upper()}%' AND permittype='BP' AND work_class='New' AND original_zip='{zip_code}' AND issue_date >= '{two_years_ago}'"
        params = {"$where": where, "$order": "issue_date DESC", "$limit": 100}
        for attempt in range(2):
            try:
                resp = requests.get(self.BASE_URL, params=params, timeout=20)
                if resp.status_code != 200:
                    continue
                return self._parse_permits(resp.json())
            except Exception:
                continue
        return []

    def search_street_all_types(self, street_name: str, zip_code: str) -> list[Permit]:
        """Search all permit types (building, electrical, plumbing, mechanical) for a street."""
        two_years_ago = (datetime.now() - timedelta(days=730)).strftime('%Y-%m-%dT00:00:00')
        where = f"permit_location like '%{street_name.upper()}%' AND original_zip='{zip_code}' AND issue_date >= '{two_years_ago}'"
        params = {"$where": where, "$order": "issue_date DESC", "$limit": 200}
        for attempt in range(2):
            try:
                resp = requests.get(self.BASE_URL, params=params, timeout=20)
                if resp.status_code != 200:
                    continue
                return self._parse_permits(resp.json())
            except Exception:
                continue
        return []

    def search_zip(self, zip_code: str, limit: int = 200) -> list[Permit]:
        two_years_ago = (datetime.now() - timedelta(days=730)).strftime('%Y-%m-%dT00:00:00')
        where = f"original_zip='{zip_code}' AND permittype='BP' AND work_class='New' AND issue_date >= '{two_years_ago}'"
        params = {"$where": where, "$order": "issue_date DESC", "$limit": limit}
        for attempt in range(2):
            try:
                resp = requests.get(self.BASE_URL, params=params, timeout=20)
                if resp.status_code != 200:
                    continue
                permits = self._parse_permits(resp.json())
                return [p for p in permits if 'Single Family' in p.permit_class
                        or 'Two Family' in p.permit_class
                        or 'Secondary' in p.permit_class]
            except Exception:
                continue
        return []

    def _parse_permits(self, data: list) -> list[Permit]:
        permits = []
        for r in data:
            issue_date = ""
            if r.get("issue_date"):
                try:
                    dt = datetime.fromisoformat(r["issue_date"].replace("T", " ").split(".")[0])
                    issue_date = dt.strftime("%Y-%m-%d")
                except Exception:
                    issue_date = str(r["issue_date"])[:10]
            permits.append(Permit(
                address=r.get("permit_location", ""),
                description=r.get("description", ""),
                sqft=float(r.get("total_new_add_sqft", 0) or 0),
                issue_date=issue_date,
                permit_class=r.get("permit_class", ""),
                work_class=r.get("work_class", ""),
                builder=r.get("contractor_company_name", "Unknown"),
                contractor_name=r.get("contractor_full_name", ""),
                applicant_name=r.get("applicant_full_name", ""),
                applicant_org=r.get("applicant_org", ""),
                housing_units=int(r.get("housing_units", 0) or 0),
                floors=int(r.get("number_of_floors", 0) or 0),
                status=r.get("status_current", ""),
                permit_number=r.get("permit_num", r.get("permitnumber", "")),
                permit_type=r.get("permittype", ""),
                permit_type_desc=r.get("permit_type_desc", ""),
            ))
        return permits


# ── Redfin Comps Module ──

class RedfinComps:
    REGION_CACHE = {}

    def _get_region_id(self, zip_code: str) -> Optional[str]:
        if zip_code in self.REGION_CACHE:
            return self.REGION_CACHE[zip_code]
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36'
        }
        try:
            resp = requests.get(f'https://www.redfin.com/zipcode/{zip_code}', headers=headers, timeout=15)
            if resp.status_code == 200:
                matches = re.findall(r'region_id=(\d+)', resp.text)
                for rid in matches:
                    if rid != zip_code and len(rid) >= 4:
                        self.REGION_CACHE[zip_code] = rid
                        return rid
                match2 = re.search(r'"regionId"\s*:\s*(\d+)', resp.text)
                if match2 and match2.group(1) != zip_code:
                    rid = match2.group(1)
                    self.REGION_CACHE[zip_code] = rid
                    return rid
        except Exception:
            pass
        return None

    def get_sold_comps(self, zip_code: str, lat: float = 0, lon: float = 0, radius_miles: float = 0) -> list[dict]:
        """Get sold comps in ZIP via region_id. If lat/lon/radius provided, compute distance and filter."""
        region_id = self._get_region_id(zip_code)
        if not region_id:
            return []
        user_agents = [
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 14_5) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.4 Safari/605.1.15',
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:128.0) Gecko/20100101 Firefox/128.0',
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36',
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/130.0.0.0 Safari/537.36',
        ]
        url = (
            f'https://www.redfin.com/stingray/api/gis-csv?al=1&num_homes=200'
            f'&ord=redfin-recommended-asc&page_number=1'
            f'&region_id={region_id}&region_type=2'
            f'&sold_within_days=730&status=9&uipt=1&v=8&min_year_built=2020'
        )
        for ua in user_agents:
            try:
                headers = {'User-Agent': ua}
                resp = requests.get(url, headers=headers, timeout=20)
                if resp.status_code != 200:
                    time.sleep(2)
                    continue
                lines = resp.text.strip().split('\n')
                header_idx = None
                for i, line in enumerate(lines):
                    if line.startswith('SALE TYPE') or line.startswith('"SALE TYPE'):
                        header_idx = i
                        break
                if header_idx is None:
                    continue
                reader = csv.DictReader(io.StringIO('\n'.join(lines[header_idx:])))
                comps = []
                for row in reader:
                    try:
                        price_str = (row.get('PRICE') or '0').replace(',', '').replace('$', '')
                        price = int(float(price_str)) if price_str else 0
                        sqft_str = (row.get('SQUARE FEET') or '0').replace(',', '')
                        sqft = int(float(sqft_str)) if sqft_str else 0
                        year_str = row.get('YEAR BUILT') or '0'
                        year = int(float(year_str)) if year_str else 0
                        psf = round(price / sqft) if sqft > 0 else 0
                        if price > 0 and year >= 2020:
                            redfin_url = ''
                            for key in row.keys():
                                if key and 'URL' in key.upper():
                                    redfin_url = row[key] or ''
                                    break
                            raw_addr = (row.get('ADDRESS') or '').strip()
                            city = (row.get('CITY') or '').strip()
                            state = (row.get('STATE OR PROVINCE') or 'TX').strip()
                            zipcode = (row.get('ZIP OR POSTAL CODE') or '').strip()
                            zillow_query = f"{raw_addr} {city} {state} {zipcode}".replace(' ', '-')
                            dist = 0
                            comp_lat = float(row.get('LATITUDE') or 0)
                            comp_lon = float(row.get('LONGITUDE') or 0)
                            if comp_lat and comp_lon and lat and lon:
                                dist = ((comp_lat - lat) * 69) ** 2 + ((comp_lon - lon) * 60) ** 2
                                dist = dist ** 0.5
                            comps.append({
                                'address': f"{raw_addr}, {city}",
                                'price': price, 'sqft': sqft, 'psf': psf,
                                'year_built': year,
                                'sold_date': row.get('SOLD DATE') or '',
                                'beds': row.get('BEDS') or '',
                                'baths': row.get('BATHS') or '',
                                'redfin_url': redfin_url,
                                'zillow_url': f"https://www.zillow.com/homes/{zillow_query}_rb/",
                                'distance_mi': round(dist, 2),
                            })
                    except (ValueError, ZeroDivisionError):
                        continue
                if radius_miles > 0 and lat and lon:
                    comps = [c for c in comps if c.get('distance_mi', 99) <= radius_miles]
                return sorted(comps, key=lambda x: x.get('distance_mi', 99))
            except Exception:
                time.sleep(2)
                continue
        return []

    def get_active_listings(self, zip_code: str, lat: float = 0, lon: float = 0, radius_miles: float = 1.0) -> list[dict]:
        """Get currently active (for sale) listings in ZIP, filtered by distance if lat/lon provided."""
        region_id = self._get_region_id(zip_code)
        if not region_id:
            return []
        user_agents = [
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 14_5) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.4 Safari/605.1.15',
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:128.0) Gecko/20100101 Firefox/128.0',
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36',
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/130.0.0.0 Safari/537.36',
        ]
        url = (
            f'https://www.redfin.com/stingray/api/gis-csv?al=1&num_homes=200'
            f'&ord=redfin-recommended-asc&page_number=1'
            f'&region_id={region_id}&region_type=2'
            f'&status=1&uipt=1&v=8'
        )
        for ua in user_agents:
            try:
                headers = {'User-Agent': ua}
                resp = requests.get(url, headers=headers, timeout=20)
                if resp.status_code != 200:
                    time.sleep(2)
                    continue
                lines = resp.text.strip().split('\n')
                header_idx = None
                for i, line in enumerate(lines):
                    if line.startswith('SALE TYPE') or line.startswith('"SALE TYPE'):
                        header_idx = i
                        break
                if header_idx is None:
                    continue
                reader = csv.DictReader(io.StringIO('\n'.join(lines[header_idx:])))
                comps = []
                for row in reader:
                    try:
                        price_str = (row.get('PRICE') or '0').replace(',', '').replace('$', '')
                        price = int(float(price_str)) if price_str else 0
                        sqft_str = (row.get('SQUARE FEET') or '0').replace(',', '')
                        sqft = int(float(sqft_str)) if sqft_str else 0
                        psf = round(price / sqft) if sqft > 0 else 0
                        if price > 0:
                            redfin_url = ''
                            for key in row.keys():
                                if key and 'URL' in key.upper():
                                    redfin_url = row[key] or ''
                                    break
                            raw_addr = (row.get('ADDRESS') or '').strip()
                            city = (row.get('CITY') or '').strip()
                            state = (row.get('STATE OR PROVINCE') or 'TX').strip()
                            zipcode = (row.get('ZIP OR POSTAL CODE') or '').strip()
                            zillow_query = f"{raw_addr} {city} {state} {zipcode}".replace(' ', '-')
                            comp_lat = float(row.get('LATITUDE') or 0)
                            comp_lon = float(row.get('LONGITUDE') or 0)
                            dist = 0
                            if comp_lat and comp_lon:
                                dist = ((comp_lat - lat) * 69) ** 2 + ((comp_lon - lon) * 60) ** 2
                                dist = dist ** 0.5
                            comps.append({
                                'address': f"{raw_addr}, {city}",
                                'price': price, 'sqft': sqft, 'psf': psf,
                                'year_built': int(float(row.get('YEAR BUILT') or 0)),
                                'beds': row.get('BEDS') or '',
                                'baths': row.get('BATHS') or '',
                                'days_on_market': row.get('DAYS ON MARKET') or '',
                                'redfin_url': redfin_url,
                                'zillow_url': f"https://www.zillow.com/homes/{zillow_query}_rb/",
                                'distance_mi': round(dist, 2),
                            })
                    except (ValueError, ZeroDivisionError):
                        continue
                # Filter to within radius if lat/lon provided
                if lat and lon:
                    comps = [c for c in comps if c.get('distance_mi', 99) <= radius_miles]
                return sorted(comps, key=lambda x: x.get('distance_mi', 99))
            except Exception:
                time.sleep(2)
                continue
        return []

    def get_rental_listings(self, zip_code: str, lat: float = 0, lon: float = 0, radius_miles: float = 2.0) -> list[dict]:
        """Get active rental listings in ZIP from Redfin, filtered by distance if lat/lon provided."""
        region_id = self._get_region_id(zip_code)
        if not region_id:
            return []
        user_agents = [
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36',
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/130.0.0.0 Safari/537.36',
        ]
        # Redfin rental CSV: status=1 (active), is_rentals=true
        url = (
            f'https://www.redfin.com/stingray/api/gis-csv?al=1&num_homes=200'
            f'&ord=redfin-recommended-asc&page_number=1'
            f'&region_id={region_id}&region_type=2'
            f'&status=1&uipt=1,2,3&v=8&is_rentals=true'
        )
        for ua in user_agents:
            try:
                headers = {'User-Agent': ua}
                resp = requests.get(url, headers=headers, timeout=20)
                if resp.status_code != 200:
                    time.sleep(2)
                    continue
                lines = resp.text.strip().split('\n')
                header_idx = None
                for i, line in enumerate(lines):
                    if line.startswith('SALE TYPE') or line.startswith('"SALE TYPE') or 'PRICE' in line.upper()[:50]:
                        header_idx = i
                        break
                if header_idx is None:
                    continue
                reader = csv.DictReader(io.StringIO('\n'.join(lines[header_idx:])))
                rentals = []
                for row in reader:
                    try:
                        # For rentals, PRICE is monthly rent
                        price_str = (row.get('PRICE') or row.get('PRICE/SQ.FT.') or '0').replace(',', '').replace('$', '').replace('/mo', '')
                        rent = int(float(price_str)) if price_str else 0
                        sqft_str = (row.get('SQUARE FEET') or '0').replace(',', '')
                        sqft = int(float(sqft_str)) if sqft_str else 0
                        rent_psf = round(rent / sqft, 2) if sqft > 0 else 0
                        if rent > 0:
                            redfin_url = ''
                            for key in row.keys():
                                if key and 'URL' in key.upper():
                                    redfin_url = row[key] or ''
                                    break
                            raw_addr = (row.get('ADDRESS') or '').strip()
                            city = (row.get('CITY') or '').strip()
                            comp_lat = float(row.get('LATITUDE') or 0)
                            comp_lon = float(row.get('LONGITUDE') or 0)
                            dist = 0
                            if comp_lat and comp_lon and lat and lon:
                                dist = ((comp_lat - lat) * 69) ** 2 + ((comp_lon - lon) * 60) ** 2
                                dist = dist ** 0.5
                            rentals.append({
                                'address': f"{raw_addr}, {city}",
                                'rent': rent, 'sqft': sqft, 'rent_psf': rent_psf,
                                'beds': row.get('BEDS') or '',
                                'baths': row.get('BATHS') or '',
                                'property_type': row.get('PROPERTY TYPE') or row.get('HOME TYPE') or '',
                                'redfin_url': redfin_url,
                                'distance_mi': round(dist, 2),
                            })
                    except (ValueError, ZeroDivisionError):
                        continue
                if lat and lon:
                    rentals = [r for r in rentals if r.get('distance_mi', 99) <= radius_miles]
                return sorted(rentals, key=lambda x: x.get('distance_mi', 99))
            except Exception:
                time.sleep(2)
                continue
        return []

    def check_listing_status(self, address: str, zip_code: str) -> dict:
        """Check if property is active, pending, or sold on Redfin."""
        region_id = self._get_region_id(zip_code)
        if not region_id:
            return {'status': 'Unknown', 'url': ''}
        user_agents = [
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 14_5) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.4 Safari/605.1.15',
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:128.0) Gecko/20100101 Firefox/128.0',
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36',
        ]
        # Extract street number + name for matching
        addr_parts = address.upper().replace(',', '').split()
        addr_num = addr_parts[0] if addr_parts else ''
        addr_street = addr_parts[1] if len(addr_parts) > 1 else ''

        # Check: 1=Active, 130=Pending/Under Contract, 9=Sold
        for status_code, label in [(1, 'Active'), (130, 'Pending'), (9, 'Sold')]:
            for ua in user_agents:
                try:
                    url = (
                        f'https://www.redfin.com/stingray/api/gis-csv?al=1&num_homes=100'
                        f'&region_id={region_id}&region_type=2'
                        f'&status={status_code}&uipt=1&v=8'
                    )
                    headers = {'User-Agent': ua}
                    resp = requests.get(url, headers=headers, timeout=20)
                    if resp.status_code != 200:
                        continue
                    lines = resp.text.strip().split('\n')
                    for i, line in enumerate(lines):
                        if 'SALE TYPE' in line:
                            reader = csv.DictReader(io.StringIO('\n'.join(lines[i:])))
                            for row in reader:
                                row_addr = (row.get('ADDRESS') or '').upper()
                                if addr_num in row_addr and addr_street in row_addr:
                                    redfin_url = ''
                                    for key in row.keys():
                                        if key and 'URL' in key.upper():
                                            redfin_url = row[key] or ''
                                            break
                                    return {
                                        'status': label,
                                        'price': row.get('PRICE', ''),
                                        'url': redfin_url,
                                    }
                            break
                except Exception:
                    continue
        return {'status': 'Not Found', 'url': ''}

def generate_report_bytes(result: AnalysisResult, address: str, zip_code: str,
                          street_name: str, purchase_price: float,
                          build_sf: float, exit_psf: float, build_cost_psf: float) -> bytes:
    """Generate Word report and return as bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    def set_cell_shading(cell, color):
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    def add_table(headers, rows, header_color='1F4E79'):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_shading(cell, header_color)
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = str(val)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)
                if r_idx % 2 == 1:
                    set_cell_shading(cell, 'F2F2F2')
        return table

    def fmt(val):
        return f"${val:,.0f}" if val else "N/A"

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f'{address}, {zip_code}')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Investment Analysis Report')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(89, 89, 89)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(datetime.now().strftime('%B %Y'))
    doc.add_page_break()

    # Permits section
    doc.add_heading(f'Construction Permits — {street_name} St', level=1)
    if result.street_permits:
        active = [p for p in result.street_permits if p.status == 'Active']
        final = [p for p in result.street_permits if p.status == 'Final']
        if active:
            doc.add_heading('Under Construction (Active)', level=2)
            rows = [[p.address, f"{p.sqft:,.0f} sf", p.issue_date, p.builder, p.permit_class]
                    for p in active]
            add_table(['Address', 'Size', 'Permit Date', 'Builder', 'Type'], rows, header_color='D4A017')
        if final:
            doc.add_heading('Completed Projects', level=2)
            rows = [[p.address, f"{p.sqft:,.0f} sf", p.issue_date, p.builder, p.permit_class]
                    for p in final]
            add_table(['Address', 'Size', 'Permit Date', 'Builder', 'Type'], rows)
    else:
        doc.add_paragraph('No new construction permits found.')

    doc.add_page_break()

    # Redfin comps section
    doc.add_heading(f'Sold Comps — {zip_code} (New Construction)', level=1)
    if result.redfin_comps:
        stats = result.market_stats
        add_table(['Metric', 'Value'], [
            ['Median $/sf', f"${stats.get('median_psf', 0)}/sf"],
            ['Average $/sf', f"${stats.get('avg_psf', 0)}/sf"],
            ['Range', f"${stats.get('min_psf', 0)} — ${stats.get('max_psf', 0)}/sf"],
            ['Count', str(stats.get('count', 0))],
        ])
        doc.add_paragraph()
        comps_sorted = sorted(result.redfin_comps, key=lambda x: x.get("psf", 0), reverse=True)
        rows = [[c["address"], fmt(c.get("price", 0)), f"{c.get('sqft', 0):,} sf",
                 f"${c.get('psf', 0)}/sf", str(c.get('year_built', '')), c.get('sold_date', '')]
                for c in comps_sorted if c.get("psf", 0) > 0]
        add_table(['Address', 'Price', 'Size', '$/sf', 'Built', 'Sold'], rows)
    else:
        doc.add_paragraph('No Redfin comps available.')

    doc.add_page_break()

    # Investment analysis
    if purchase_price > 0 and build_sf > 0:
        doc.add_heading('Investment Analysis', level=1)
        total_build_cost = build_cost_psf * build_sf
        median_psf = result.market_stats.get("median_psf", 0)

        add_table(['Parameter', 'Value'], [
            ['Purchase Price', fmt(purchase_price)],
            ['Build Size', f"{build_sf:,.0f} sf"],
            [f'Build Cost @ ${build_cost_psf:,.0f}/sf', fmt(total_build_cost)],
            ['Exit Assumption', f"${exit_psf}/sf" if exit_psf else "N/A"],
            ['Market Median', f"${median_psf}/sf" if median_psf else "N/A"],
        ])

        doc.add_heading('Exit Scenarios', level=2)
        scenarios = []
        test_psfs = set()
        if median_psf:
            test_psfs.update([median_psf, median_psf + 25, median_psf + 50])
        if exit_psf:
            test_psfs.add(exit_psf)
        for psf in sorted(test_psfs):
            if psf > 0:
                revenue = psf * build_sf
                total_cost = purchase_price + total_build_cost + (total_build_cost * 0.10)
                profit = revenue - total_cost
                margin = (profit / total_cost) * 100
                scenarios.append([f"${psf}/sf", fmt(revenue), fmt(total_cost), fmt(profit),
                                  f"{margin:.1f}%", "✅" if margin > 10 else ("⚠" if margin > 0 else "❌")])
        add_table(['Exit $/sf', 'Revenue', 'Total Cost', 'Profit', 'Margin', ''], scenarios)

    # Sources
    doc.add_page_break()
    doc.add_heading('Data Sources', level=1)
    for name, desc, status in [
        ('Austin Open Data', 'Construction Permits API', result.sources_status.get('permits', '✅')),
        ('Redfin', 'Sold comps CSV API', result.sources_status.get('redfin', '✅')),
        ('TCAD', 'Property appraisals & deeds', '⚠ Run locally with CLI tool for TCAD data'),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f'{status} {name}: ')
        run.bold = True
        p.add_run(desc)

    p = doc.add_paragraph()
    run = p.add_run('\nDisclaimer: ')
    run.bold = True
    p.add_run('Texas is a non-disclosure state. Actual sale prices are not in public deed records. '
              'This report is for informational purposes only.')

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Analysis Runner ──

@st.cache_data(ttl=3600, show_spinner=False)
def run_analysis(address: str, zip_code: str, street_name: str):
    """Run permits + Redfin analysis (cached for 1 hour)."""
    result = AnalysisResult()
    permits_api = AustinPermits()
    redfin_api = RedfinComps()

    # Permits
    result.street_permits = permits_api.search_street(street_name, zip_code)
    result.street_all_permits = permits_api.search_street_all_types(street_name, zip_code)
    result.zip_permits = permits_api.search_zip(zip_code)
    result.sources_status['permits'] = '✅' if result.street_permits or result.zip_permits else '⚠'

    # Geocode for radius-based queries
    lat, lon = geocode_address(address, zip_code)

    # Redfin — Sold comps (1 mile radius, fallback to ZIP-wide)
    comps_radius = True
    if lat and lon:
        result.redfin_comps = redfin_api.get_sold_comps(zip_code, lat, lon, radius_miles=1.0)
    if not result.redfin_comps:
        result.redfin_comps = redfin_api.get_sold_comps(zip_code, lat, lon, radius_miles=0)
        comps_radius = False
    if result.redfin_comps:
        psf_values = sorted([c["psf"] for c in result.redfin_comps if c.get("psf", 0) > 0])
        if psf_values:
            mid = len(psf_values) // 2
            result.market_stats = {
                "median_psf": psf_values[mid],
                "avg_psf": round(sum(psf_values) / len(psf_values)),
                "min_psf": min(psf_values),
                "max_psf": max(psf_values),
                "count": len(psf_values),
            }
        result.sources_status['redfin'] = '✅'
        result.sources_status['comps_radius'] = comps_radius
    else:
        result.sources_status['redfin'] = '⚠ No data'
        # Fallback: use Census Bureau median home value to estimate $/sf
        try:
            census = fetch_census_home_value(zip_code)
            if census.get('median_value') and census['median_value'] > 0:
                # Estimate $/sf from Census median value / typical new construction size
                typical_sf = 1800  # typical new construction in Austin
                est_psf = round(census['median_value'] / typical_sf)
                result.market_stats = {
                    "median_psf": est_psf,
                    "avg_psf": est_psf,
                    "min_psf": est_psf,
                    "max_psf": est_psf,
                    "count": 0,
                    "source": "Census Bureau (estimated)",
                }
                result.sources_status['redfin'] = '⚠ Census fallback'
        except Exception:
            pass

    # Listing status (active/pending/sold)
    result.listing_status = redfin_api.check_listing_status(address, zip_code)

    # Neighborhood comps are now the same as sold comps (both 1 mile)
    if lat and lon:
        result.neighborhood_comps = result.redfin_comps
        result.neighborhood_stats = result.market_stats

        # Active listings (1 mile radius)
        result.active_comps = redfin_api.get_active_listings(zip_code, lat, lon, radius_miles=1.0)
        if result.active_comps:
            a_psf = sorted([c["psf"] for c in result.active_comps if c.get("psf", 0) > 0])
            if a_psf:
                mid = len(a_psf) // 2
                result.active_stats = {
                    "median_psf": a_psf[mid],
                    "avg_psf": round(sum(a_psf) / len(a_psf)),
                    "min_psf": min(a_psf),
                    "max_psf": max(a_psf),
                    "count": len(a_psf),
                }

        # Rental listings — Redfin CSV doesn't support rentals, so skip
        # Users can research rentals via the links in the Rental Comps tab

    return result


def extract_street_name(address: str) -> str:
    parts = address.upper().replace(",", "").split()
    suffixes = {"ST", "STREET", "AVE", "AVENUE", "DR", "DRIVE", "LN", "LANE",
                "BLVD", "BOULEVARD", "CT", "COURT", "WAY", "RD", "ROAD",
                "CIR", "CIRCLE", "PL", "PLACE", "TRL", "TRAIL", "PKWY", "PARKWAY"}
    street_parts = [p for p in parts[1:] if p not in suffixes]
    return " ".join(street_parts) if street_parts else (parts[1] if len(parts) > 1 else parts[0])


# ── Plot Info Module ──

ARCGIS_BASE = "https://services.arcgis.com/0L95CJ0VTaxqcmED/ArcGIS/rest/services"

# Austin zoning density rules (approximate max units per lot)
ZONING_INFO = {
    "SF-1": {
        "desc": "Single Family Residence - Large Lot",
        "max_units": 1, "min_lot_sf": 10000,
        "plain": "Only one house allowed. Large lot (¼ acre+). Think suburban estate feel.",
        "can_build": "1 single-family home + 1 ADU (accessory dwelling unit, like a garage apartment)",
        "height": "35 ft (2-3 stories)",
    },
    "SF-2": {
        "desc": "Single Family Residence - Standard Lot",
        "max_units": 1, "min_lot_sf": 5750,
        "plain": "Typical Austin residential neighborhood. One house per lot, standard-sized yard.",
        "can_build": "1 single-family home + 1 ADU. Duplex NOT allowed unless you get a zoning change.",
        "height": "35 ft (2-3 stories)",
    },
    "SF-3": {
        "desc": "Single Family Residence - Standard Lot (more flexible)",
        "max_units": 1, "min_lot_sf": 5750,
        "plain": "Same as SF-2 but slightly more flexible. Most common residential zoning in Austin.",
        "can_build": "1 single-family home + 1 ADU. Duplexes allowed on corner lots in some cases.",
        "height": "35 ft (2-3 stories)",
    },
    "SF-4A": {
        "desc": "Single Family - Small Lot",
        "max_units": 1, "min_lot_sf": 3500,
        "plain": "Smaller lots, urban infill. Great for compact new construction.",
        "can_build": "1 single-family home + 1 ADU on a smaller lot",
        "height": "35 ft (2-3 stories)",
    },
    "SF-5": {
        "desc": "Single Family - Urban",
        "max_units": 1, "min_lot_sf": 2500,
        "plain": "Very small urban lots. Townhome-style development possible.",
        "can_build": "1 single-family home or townhome + 1 ADU",
        "height": "35 ft (2-3 stories)",
    },
    "SF-6": {
        "desc": "Townhouse / Condo",
        "max_units": 8, "min_lot_sf": 2500,
        "plain": "Allows multiple attached units (townhomes, condos). Good for small-scale development.",
        "can_build": "Up to 8 townhome/condo units depending on lot size",
        "height": "35 ft (2-3 stories)",
    },
    "MF-1": {
        "desc": "Multifamily - Low Density",
        "max_units": "18/acre", "min_lot_sf": 8000,
        "plain": "Small apartment buildings, duplexes, fourplexes. Residential feel but multiple units.",
        "can_build": "~18 units per acre. On a 7,000 sf lot ≈ 2-3 units.",
        "height": "40 ft (3 stories)",
    },
    "MF-2": {
        "desc": "Multifamily - Low-Medium Density",
        "max_units": "25/acre", "min_lot_sf": 8000,
        "plain": "Medium apartment buildings. Common along transit corridors.",
        "can_build": "~25 units per acre. On a 7,000 sf lot ≈ 4 units.",
        "height": "40 ft (3 stories)",
    },
    "MF-3": {
        "desc": "Multifamily - Medium Density",
        "max_units": "36/acre", "min_lot_sf": 8000,
        "plain": "Larger apartment complexes. Urban mixed-use areas.",
        "can_build": "~36 units per acre.",
        "height": "40 ft (3 stories)",
    },
    "MF-4": {
        "desc": "Multifamily - Moderate-High Density",
        "max_units": "54/acre", "min_lot_sf": 8000,
        "plain": "Dense apartment buildings. Downtown-adjacent areas.",
        "can_build": "~54 units per acre.",
        "height": "60 ft (5 stories)",
    },
    "MF-5": {
        "desc": "Multifamily - High Density",
        "max_units": "No max", "min_lot_sf": 8000,
        "plain": "High-rise apartments. No unit cap — limited by building size/FAR.",
        "can_build": "No unit maximum. Limited by floor-area ratio and height.",
        "height": "60 ft (5 stories)",
    },
    "MF-6": {
        "desc": "Multifamily - Highest Density",
        "max_units": "No max", "min_lot_sf": 10000,
        "plain": "Tallest residential buildings. Downtown high-rises.",
        "can_build": "No unit maximum. Tallest allowed residential.",
        "height": "No limit",
    },
}

# Plain-English overlay explanations
OVERLAY_EXPLANATIONS = {
    "-NP": (
        "🏘️ **Neighborhood Plan (NP)**",
        "This property is in a Neighborhood Plan area. The neighborhood has agreed to extra rules about "
        "what can be built — things like building height, setbacks, parking, and design standards. "
        "You may need to attend a neighborhood meeting before getting permits. "
        "Check the specific plan at [Austin Neighborhood Plans](https://www.austintexas.gov/department/neighborhood-plans)."
    ),
    "-CO": (
        "📋 **Conditional Overlay (CO)**",
        "This lot has special conditions attached by the City. These are specific rules that override "
        "the base zoning — for example, limiting hours of operation, requiring extra landscaping, or "
        "restricting certain uses. You MUST check the specific conditions with the City of Austin."
    ),
    "-H": (
        "🏛️ **Historic (H)**",
        "This property is in a historic district. Renovations and new construction must follow strict "
        "design guidelines to preserve neighborhood character. Demolition may be restricted or prohibited."
    ),
    "-V": (
        "🌿 **Vertical Mixed Use (V/VMU)**",
        "Allows ground-floor commercial with residential above. Great for mixed-use development. "
        "May get density bonuses if affordable housing is included."
    ),
}


@st.cache_data(ttl=3600)
def fetch_census_rents(zip_code: str) -> dict:
    """Fetch median rent data from Census Bureau ACS 5-year survey (no API key needed)."""
    try:
        url = 'https://api.census.gov/data/2022/acs/acs5'
        params = {
            'get': 'B25064_001E,B25031_002E,B25031_003E,B25031_004E,B25031_005E,B25031_006E',
            'for': f'zip code tabulation area:{zip_code}'
        }
        resp = requests.get(url, params=params, timeout=15)
        if resp.status_code != 200:
            return {}
        data = resp.json()
        if len(data) < 2:
            return {}
        values = data[1]
        def _int(v):
            try:
                return int(v) if v and int(v) > 0 else None
            except (ValueError, TypeError):
                return None
        return {
            'median_rent': _int(values[0]),
            'rent_0br': _int(values[1]),
            'rent_1br': _int(values[2]),
            'rent_2br': _int(values[3]),
            'rent_3br': _int(values[4]),
            'rent_4br_plus': _int(values[5]),
            'source': 'Census ACS 5-Year (2022)',
        }
    except Exception:
        return {}


def fetch_census_home_value(zip_code: str) -> dict:
    """Fetch median home value from Census Bureau ACS (no API key needed)."""
    try:
        url = 'https://api.census.gov/data/2022/acs/acs5'
        params = {
            'get': 'B25077_001E',
            'for': f'zip code tabulation area:{zip_code}'
        }
        resp = requests.get(url, params=params, timeout=15)
        if resp.status_code != 200:
            return {}
        data = resp.json()
        if len(data) < 2:
            return {}
        val = data[1][0]
        return {'median_value': int(val) if val else 0}
    except Exception:
        return {}


def geocode_address(address: str, zip_code: str):
    """Geocode address using Census Bureau geocoder, fallback to Nominatim."""
    # Try Census geocoder first
    try:
        params = {
            'address': f'{address}, Austin, TX {zip_code}',
            'benchmark': 'Public_AR_Current',
            'format': 'json',
        }
        r = requests.get('https://geocoding.geo.census.gov/geocoder/locations/onelineaddress',
                        params=params, timeout=15)
        matches = r.json().get('result', {}).get('addressMatches', [])
        if matches:
            coords = matches[0]['coordinates']
            return coords['y'], coords['x']
    except Exception:
        pass
    # Fallback to Nominatim
    try:
        r = requests.get('https://nominatim.openstreetmap.org/search',
                        params={'q': f'{address}, Austin, TX {zip_code}', 'format': 'json', 'limit': 1},
                        headers={'User-Agent': 'RE-Analyzer/1.0'}, timeout=15)
        data = r.json()
        if data:
            return float(data[0]['lat']), float(data[0]['lon'])
    except Exception:
        pass
    return None, None


@st.cache_data(ttl=3600)
def fetch_plot_info(lat: float, lon: float, address: str = ""):
    """Fetch zoning, parcel, and flood data from Austin ArcGIS."""
    plot_data = {}
    buf = 0.0003  # ~30 meters buffer for envelope queries

    # Extract address number for parcel matching
    addr_num = ""
    parts = address.split()
    if parts and parts[0].isdigit():
        addr_num = parts[0]

    # Zoning (use envelope/buffer since point can miss on parcel boundaries)
    try:
        url = f'{ARCGIS_BASE}/Current_Zoning_gdb/FeatureServer/0/query'
        buf = 0.0003  # ~30 meters buffer
        params = {
            'geometry': f'{lon-buf},{lat-buf},{lon+buf},{lat+buf}',
            'geometryType': 'esriGeometryEnvelope',
            'spatialRel': 'esriSpatialRelIntersects',
            'inSR': '4326', 'outFields': '*', 'f': 'json', 'returnGeometry': 'false',
        }
        r = requests.get(url, params=params, timeout=15)
        features = r.json().get('features', [])
        if features:
            attrs = features[0]['attributes']
            plot_data['zoning'] = {
                'zoning_type': attrs.get('ZONING_ZTYPE', ''),
                'base_zone': attrs.get('BASE_ZONE', ''),
                'zone_name': attrs.get('ZONE_NAME', ''),
                'lot_area_sf': round(attrs.get('SHAPE__Area', 0)),
            }
    except Exception:
        pass

    # TCAD Parcel (use envelope since point can miss on boundaries)
    try:
        url = f'{ARCGIS_BASE}/EXTERNAL_tcad_parcel/FeatureServer/0/query'
        params = {
            'geometry': f'{lon-buf},{lat-buf},{lon+buf},{lat+buf}',
            'geometryType': 'esriGeometryEnvelope',
            'spatialRel': 'esriSpatialRelIntersects',
            'inSR': '4326', 'outFields': '*', 'f': 'json', 'returnGeometry': 'false',
        }
        r = requests.get(url, params=params, timeout=15)
        features = r.json().get('features', [])
        if features:
            # Match by address number if multiple parcels returned
            best = features[0]
            if addr_num and len(features) > 1:
                for f in features:
                    situs = str(f['attributes'].get('SITUS', ''))
                    if situs == addr_num:
                        best = f
                        break
            attrs = best['attributes']
            plot_data['parcel'] = {
                'prop_id': attrs.get('PROP_ID', ''),
                'pid': attrs.get('PID_10', ''),
                'situs': attrs.get('SITUS', ''),
                'lot': attrs.get('LOTS', ''),
                'block': attrs.get('BLOCKS', ''),
                'parcel_area_sf': round(attrs.get('Shape__Area', 0)),
            }
    except Exception:
        pass

    # FEMA Flood (use envelope)
    try:
        url = f'{ARCGIS_BASE}/INLANDWATERS_greater_austin_fema_floodplain/FeatureServer/0/query'
        params = {
            'geometry': f'{lon-buf},{lat-buf},{lon+buf},{lat+buf}',
            'geometryType': 'esriGeometryEnvelope',
            'spatialRel': 'esriSpatialRelIntersects',
            'inSR': '4326', 'outFields': '*', 'f': 'json', 'returnGeometry': 'false',
        }
        r = requests.get(url, params=params, timeout=15)
        features = r.json().get('features', [])
        plot_data['flood'] = {
            'in_floodplain': len(features) > 0,
            'zone': features[0]['attributes'].get('FLD_ZONE', '') if features else 'None (X - Minimal Risk)',
        }
    except Exception:
        plot_data['flood'] = {'in_floodplain': False, 'zone': 'Unable to determine'}

    return plot_data


# ══════════════════════════════════════════════════════════════
#  EXCEL EXPORT — Generate multi-sheet workbook
# ══════════════════════════════════════════════════════════════

def _compute_hold_profit(purchase_price, build_sf, build_cost, exit_price, exit_cost_pct,
                         hard_contingency_pct, soft_cost_pct, soft_contingency,
                         ltv, interest_rate, draw_factor, loan_fee_pct,
                         build_months, hold_months, delay_months,
                         perm_mortgage_rate, amortization_years,
                         rent_per_unit, units, vacancy_pct, mgmt_fee_pct,
                         prop_tax_rate, taxable_value_psf, insurance_monthly,
                         repairs_per_unit, common_utilities, leasing_reserve):
    """Recompute profit & equity multiple for a given build_cost/exit_price combo."""
    hc = build_cost * build_sf
    hcont = hc * (hard_contingency_pct / 100)
    sc = hc * (soft_cost_pct / 100)
    tdc = hc + hcont + sc + soft_contingency
    tpc = purchase_price + tdc

    la = tdc * (ltv / 100)
    ci = la * (interest_rate / 100) * (draw_factor / 100) * (build_months + delay_months) / 12
    lf = la * (loan_fee_pct / 100)

    if hold_months > 0:
        mpr = perm_mortgage_rate / 100 / 12
        npay = amortization_years * 12
        if mpr > 0:
            mds = la * (mpr * (1 + mpr) ** npay) / ((1 + mpr) ** npay - 1)
        else:
            mds = la / npay
        hi = mds * hold_months

        if mpr > 0:
            lbal = la * (1 + mpr) ** hold_months - mds * ((1 + mpr) ** hold_months - 1) / mpr
        else:
            lbal = la - mds * hold_months

        gr = rent_per_unit * units * hold_months
        er = gr * (1 - vacancy_pct / 100)
        mc = er * (mgmt_fee_pct / 100)
        pt = build_sf * taxable_value_psf * (prop_tax_rate / 100) * hold_months / 12
        ins = insurance_monthly * hold_months
        rep = repairs_per_unit * units * hold_months
        mi = common_utilities * hold_months
        lea = leasing_reserve * hold_months
        mnoi = (er - mc - pt - ins - rep - mi - lea) / max(hold_months, 1)
        mcf = mnoi - mds
        cum_cf = mcf * hold_months
        aeq = max(0, -cum_cf)
        tei = purchase_price + ci + lf + aeq

        rev = exit_price * build_sf
        ec = rev * (exit_cost_pct / 100)
        nsbd = rev - ec
        nsad = nsbd - lbal
        pcf = max(0, cum_cf)
        tcr = nsad + pcf
        profit = tcr - tei
        em = tcr / tei if tei > 0 else 0
    else:
        tc = tpc + ci + lf + (exit_price * build_sf * exit_cost_pct / 100)
        tei = tc
        rev = exit_price * build_sf
        profit = rev - tc
        em = rev / tc if tc > 0 else 0

    return profit, em


def generate_excel_bytes(
    # Sidebar inputs
    units, per_unit_sf, build_sf, purchase_price, build_cost_psf, exit_psf,
    build_months, hold_months, delay_months,
    hard_contingency_pct, soft_cost_pct, soft_contingency,
    ltv, interest_rate, draw_factor, loan_fee_pct,
    perm_mortgage_rate, amortization_years,
    rent_per_unit, vacancy_pct, mgmt_fee_pct,
    prop_tax_rate, taxable_value_psf, insurance_monthly,
    repairs_per_unit, common_utilities, leasing_reserve,
    exit_cost_pct, price_decline,
    split_soft, arch_pct, eng_pct, permit_fee_pct, survey_pct, insurance_dev_pct, other_soft_pct,
    broker_fee_pct, title_closing_pct, seller_concessions_pct,
    # Computed values
    hard_cost, hard_contingency, soft_costs, total_dev_cost, total_project_cost,
    loan_amount, equity, construction_interest, loan_fees,
    hold_interest, gross_rent, effective_rent,
    total_hold_expenses, net_rental_income,
    monthly_debt_service, loan_balance_after_hold,
    total_equity_invested, user_profit, user_revenue,
    breakeven_psf, exit_costs_user, total_cost, equity_multiple,
    # Hold-specific
    additional_equity_needed, cumulative_cf, net_sale_before_debt, net_sale_after_debt,
    positive_rental_cf, total_cash_returned,
    # OPEX
    mgmt_cost, prop_tax, insurance, repairs, misc, leasing,
    monthly_noi, monthly_cf_after_debt,
):
    """Build a multi-sheet Excel workbook and return bytes."""
    wb = Workbook()

    # Styles
    bold = Font(bold=True)
    bold_white = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="4472C4")
    yellow_fill = PatternFill("solid", fgColor="FFF2CC")
    green_fill = PatternFill("solid", fgColor="E2EFDA")
    orange_fill = PatternFill("solid", fgColor="FCE4D6")
    light_blue_fill = PatternFill("solid", fgColor="D6E4F0")
    dollar_fmt = '$#,##0'
    dollar_fmt2 = '$#,##0.00'
    pct_fmt = '0.0%'
    pct_fmt2 = '0.00%'
    num_fmt = '#,##0'
    mult_fmt = '0.00x'
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    def _header_row(ws, row, cols, fill=header_fill):
        for c, val in enumerate(cols, 1):
            cell = ws.cell(row=row, column=c, value=val)
            cell.font = bold_white
            cell.fill = fill
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center')

    def _data_row(ws, row, vals, fills=None, fmts=None):
        for c, val in enumerate(vals, 1):
            cell = ws.cell(row=row, column=c, value=val)
            cell.border = thin_border
            if fills and c <= len(fills) and fills[c - 1]:
                cell.fill = fills[c - 1]
            if fmts and c <= len(fmts) and fmts[c - 1]:
                cell.number_format = fmts[c - 1]

    # ═══════════════════════════════════════════
    # SHEET 1: Hold Model
    # ═══════════════════════════════════════════
    ws1 = wb.active
    ws1.title = "Hold Model"
    ws1.sheet_properties.tabColor = "4472C4"

    # Set column widths
    ws1.column_dimensions['A'].width = 35
    ws1.column_dimensions['B'].width = 18
    ws1.column_dimensions['C'].width = 14
    ws1.column_dimensions['D'].width = 14
    ws1.column_dimensions['E'].width = 30

    r = 1
    # ── Section 1: Assumptions ──
    ws1.cell(row=r, column=1, value="ASSUMPTIONS").font = Font(bold=True, size=14)
    r += 1
    _header_row(ws1, r, ["Assumption", "Value", "Unit", "Type", "Notes"])
    r += 1

    assumptions = [
        ("Units", units, "units", "Input", ""),
        ("Sq Ft / Unit", per_unit_sf, "sf", "Calc", "build_sf / units"),
        ("Total Sq Ft", build_sf, "sf", "Input", ""),
        ("Land Equity", purchase_price, "$", "Input", "Purchase price"),
        ("Build Cost $/sf", build_cost_psf, "$/sf", "Input", ""),
        ("Exit Price $/sf", exit_psf, "$/sf", "Input", "Target sale price"),
        ("Build Duration", build_months, "months", "Input", ""),
        ("Hold Period", hold_months, "months", "Input", "After build completion"),
        ("Expected Delays", delay_months, "months", "Input", ""),
        ("Hard Cost Contingency %", hard_contingency_pct / 100, "%", "Input", ""),
        ("Soft Cost %", soft_cost_pct / 100, "%", "Input", "Of hard cost"),
        ("Soft Contingency $", soft_contingency, "$", "Input", "Fixed amount"),
        ("Construction Debt Funding %", ltv / 100, "%", "Input", "LTC on dev costs"),
        ("Construction Interest Rate %", interest_rate / 100, "%", "Input", "Annual"),
        ("Draw Factor %", draw_factor / 100, "%", "Input", ""),
        ("Loan Fees %", loan_fee_pct / 100, "%", "Input", ""),
        ("Perm Mortgage Rate %", perm_mortgage_rate / 100, "%", "Input", "Annual"),
        ("Amortization", amortization_years, "years", "Input", ""),
        ("Rent / Unit / Month", rent_per_unit, "$/mo", "Input", ""),
        ("Vacancy %", vacancy_pct / 100, "%", "Input", ""),
        ("Management Fee %", mgmt_fee_pct / 100, "%", "Input", "Of EGI"),
        ("Property Tax Rate %", prop_tax_rate / 100, "%", "Input", "Annual"),
        ("Taxable Value Basis $/sf", taxable_value_psf, "$/sf", "Input", ""),
        ("Insurance $/mo", insurance_monthly, "$/mo", "Input", ""),
        ("Repairs Reserve $/unit/mo", repairs_per_unit, "$/mo", "Input", ""),
        ("Common Utilities $/mo", common_utilities, "$/mo", "Input", ""),
        ("Leasing Reserve $/mo", leasing_reserve, "$/mo", "Input", ""),
        ("Exit Cost %", exit_cost_pct / 100, "%", "Calc", f"Broker {broker_fee_pct}% + Title {title_closing_pct}% + Concessions {seller_concessions_pct}%"),
    ]

    for label, val, unit, typ, note in assumptions:
        fmt_b = None
        if unit == "$" or unit == "$/sf" or unit == "$/mo":
            fmt_b = dollar_fmt
        elif unit == "%":
            fmt_b = pct_fmt
        _data_row(ws1, r, [label, val, unit, typ, note],
                  fills=[None, yellow_fill, None, None, None],
                  fmts=[None, fmt_b, None, None, None])
        r += 1

    # ── Section 2: Development Cost & Construction Debt ──
    r += 1
    ws1.cell(row=r, column=1, value="DEVELOPMENT COST & CONSTRUCTION DEBT").font = Font(bold=True, size=14)
    r += 1
    _header_row(ws1, r, ["Line Item", "Formula / Basis", "Amount", "Notes"])
    r += 1

    dev_items = [
        ("Hard Cost", f"{build_cost_psf} × {build_sf:,} sf", hard_cost, ""),
        ("Hard Cost Contingency", f"{hard_contingency_pct}% of Hard Cost", hard_contingency, ""),
        ("Soft Cost", f"{soft_cost_pct:.1f}% of Hard Cost", soft_costs, ""),
        ("Soft Contingency", "Fixed", soft_contingency, ""),
        ("Non-Land Development Cost", "Sum above", total_dev_cost, ""),
        ("Construction Debt", f"{ltv}% LTC", loan_amount, ""),
        ("Construction Interest During Build", f"{interest_rate}% × {draw_factor}% draw × {build_months + delay_months} mo", construction_interest, ""),
        ("Construction Loan Fees", f"{loan_fee_pct}% of loan", loan_fees, ""),
    ]
    for label, basis, amt, note in dev_items:
        _data_row(ws1, r, [label, basis, amt, note],
                  fills=[None, None, green_fill, None],
                  fmts=[None, None, dollar_fmt, None])
        r += 1

    # ── Section 3: Monthly Rental OPEX & Cash Flow ──
    r += 1
    ws1.cell(row=r, column=1, value="MONTHLY RENTAL OPEX & CASH FLOW").font = Font(bold=True, size=14)
    r += 1
    _header_row(ws1, r, ["Line Item", "Formula / Basis", "Monthly", "Annual", "Notes"])
    r += 1

    hm = max(hold_months, 1)
    monthly_gross = rent_per_unit * units
    monthly_vacancy = monthly_gross * (vacancy_pct / 100)
    monthly_egi = monthly_gross - monthly_vacancy
    monthly_mgmt = monthly_egi * (mgmt_fee_pct / 100)
    monthly_proptax = build_sf * taxable_value_psf * (prop_tax_rate / 100) / 12
    monthly_ins = insurance_monthly
    monthly_repairs = repairs_per_unit * units
    monthly_utils = common_utilities
    monthly_leasing = leasing_reserve
    monthly_total_opex = monthly_mgmt + monthly_proptax + monthly_ins + monthly_repairs + monthly_utils + monthly_leasing

    opex_items = [
        ("Gross Rent", f"{rent_per_unit:,.0f} × {units} units", monthly_gross, monthly_gross * 12, ""),
        ("Vacancy", f"{vacancy_pct}%", -monthly_vacancy, -monthly_vacancy * 12, ""),
        ("Effective Gross Income", "", monthly_egi, monthly_egi * 12, ""),
        ("Management Fee", f"{mgmt_fee_pct}% of EGI", -monthly_mgmt, -monthly_mgmt * 12, ""),
        ("Property Taxes", f"{prop_tax_rate}% of {taxable_value_psf}×{build_sf:,}", -monthly_proptax, -monthly_proptax * 12, ""),
        ("Insurance", "", -monthly_ins, -monthly_ins * 12, ""),
        ("Repairs", f"{repairs_per_unit}/unit/mo", -monthly_repairs, -monthly_repairs * 12, ""),
        ("Utilities / Misc", "", -monthly_utils, -monthly_utils * 12, ""),
        ("Leasing Reserve", "", -monthly_leasing, -monthly_leasing * 12, ""),
        ("Total OPEX", "", -monthly_total_opex, -monthly_total_opex * 12, ""),
        ("NOI", "EGI - OPEX", monthly_noi, monthly_noi * 12, ""),
        ("Permanent Debt Service", f"{perm_mortgage_rate}%, {amortization_years}yr amort", -monthly_debt_service, -monthly_debt_service * 12, ""),
        ("Monthly Cash Flow After Debt", "", monthly_cf_after_debt, monthly_cf_after_debt * 12, ""),
    ]
    for label, basis, mo, ann, note in opex_items:
        _data_row(ws1, r, [label, basis, mo, ann, note],
                  fills=[None, None, green_fill, green_fill, None],
                  fmts=[None, None, dollar_fmt, dollar_fmt, None])
        r += 1

    # ── Section 4: Profit/Loss After Hold ──
    r += 1
    ws1.cell(row=r, column=1, value="PROFIT / LOSS AFTER HOLD").font = Font(bold=True, size=14)
    r += 1
    _header_row(ws1, r, ["Line Item", "Formula / Basis", "Amount", "Notes"])
    r += 1

    profit_items = [
        ("Land Equity", "Purchase price", purchase_price, "Cash in"),
        ("Construction Interest", "During build", construction_interest, "Cash in"),
        ("Loan Fees", "", loan_fees, "Cash in"),
        ("Cumulative Rental CF", f"{hold_months} months", cumulative_cf, "Pos = cash returned"),
        ("Additional Equity for Negative CF", "max(0, -cumCF)", additional_equity_needed, "Cash in if CF negative"),
        ("Total Equity Invested", "Sum of cash in", total_equity_invested, ""),
        ("", "", "", ""),
        ("Final Sale Price", f"{exit_psf} × {build_sf:,} sf", user_revenue, ""),
        ("Exit Costs", f"{exit_cost_pct:.1f}%", -exit_costs_user, ""),
        ("Net Sale Before Debt", "", net_sale_before_debt, ""),
        ("Loan Balance After Hold", f"{hold_months} mo amort", -loan_balance_after_hold, ""),
        ("Net Sale After Debt", "", net_sale_after_debt, ""),
        ("Positive Rental CF Returned", "", positive_rental_cf, ""),
        ("Total Cash Returned", "", total_cash_returned, ""),
        ("Profit / (Loss)", "Cash returned - equity", user_profit, ""),
        ("Equity Multiple", "Cash returned / equity", equity_multiple, ""),
    ]
    for label, basis, amt, note in profit_items:
        fill_c = None
        fmt_c = dollar_fmt
        if label == "Equity Multiple":
            fmt_c = '0.00x'
            fill_c = orange_fill
        elif label == "Profit / (Loss)":
            fill_c = orange_fill
        elif label == "Total Equity Invested" or label == "Total Cash Returned":
            fill_c = light_blue_fill
        _data_row(ws1, r, [label, basis, amt, note],
                  fills=[None, None, fill_c, None],
                  fmts=[None, None, fmt_c, None])
        r += 1

    # ── Section 5: Hold Sweep — Build Cost vs Exit Price (Profit) ──
    r += 1
    ws1.cell(row=r, column=1, value="HOLD SWEEP — BUILD COST vs EXIT PRICE (Profit)").font = Font(bold=True, size=14)
    r += 1

    sweep_build_costs = [build_cost_psf - 25, build_cost_psf, build_cost_psf + 25, build_cost_psf + 50]
    sweep_exit_prices = [exit_psf - 25, exit_psf, exit_psf + 25]

    common_args = dict(
        purchase_price=purchase_price, build_sf=build_sf, exit_cost_pct=exit_cost_pct,
        hard_contingency_pct=hard_contingency_pct, soft_cost_pct=soft_cost_pct,
        soft_contingency=soft_contingency, ltv=ltv, interest_rate=interest_rate,
        draw_factor=draw_factor, loan_fee_pct=loan_fee_pct,
        build_months=build_months, hold_months=hold_months, delay_months=delay_months,
        perm_mortgage_rate=perm_mortgage_rate, amortization_years=amortization_years,
        rent_per_unit=rent_per_unit, units=units, vacancy_pct=vacancy_pct,
        mgmt_fee_pct=mgmt_fee_pct, prop_tax_rate=prop_tax_rate,
        taxable_value_psf=taxable_value_psf, insurance_monthly=insurance_monthly,
        repairs_per_unit=repairs_per_unit, common_utilities=common_utilities,
        leasing_reserve=leasing_reserve,
    )

    # Header row
    ws1.cell(row=r, column=1, value="Build Cost \\ Exit $/sf").font = bold
    ws1.cell(row=r, column=1).fill = header_fill
    ws1.cell(row=r, column=1).font = bold_white
    for j, ep in enumerate(sweep_exit_prices, 2):
        cell = ws1.cell(row=r, column=j, value=f"${ep}/sf")
        cell.font = bold_white
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
    r += 1

    for bc in sweep_build_costs:
        ws1.cell(row=r, column=1, value=f"${bc}/sf").font = bold
        for j, ep in enumerate(sweep_exit_prices, 2):
            profit_val, _ = _compute_hold_profit(build_cost=bc, exit_price=ep, **common_args)
            cell = ws1.cell(row=r, column=j, value=profit_val)
            cell.number_format = dollar_fmt
            cell.fill = green_fill if profit_val >= 0 else PatternFill("solid", fgColor="FFC7CE")
            cell.border = thin_border
        r += 1

    # ── Section 6: Equity Multiple Sweep ──
    r += 1
    ws1.cell(row=r, column=1, value="EQUITY MULTIPLE SWEEP").font = Font(bold=True, size=14)
    r += 1

    ws1.cell(row=r, column=1, value="Build Cost \\ Exit $/sf").font = bold_white
    ws1.cell(row=r, column=1).fill = header_fill
    for j, ep in enumerate(sweep_exit_prices, 2):
        cell = ws1.cell(row=r, column=j, value=f"${ep}/sf")
        cell.font = bold_white
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
    r += 1

    for bc in sweep_build_costs:
        ws1.cell(row=r, column=1, value=f"${bc}/sf").font = bold
        for j, ep in enumerate(sweep_exit_prices, 2):
            _, em_val = _compute_hold_profit(build_cost=bc, exit_price=ep, **common_args)
            cell = ws1.cell(row=r, column=j, value=em_val)
            cell.number_format = '0.00x'
            cell.fill = orange_fill if em_val >= 1 else PatternFill("solid", fgColor="FFC7CE")
            cell.border = thin_border
        r += 1

    # ═══════════════════════════════════════════
    # SHEET 2: Detailed Analysis
    # ═══════════════════════════════════════════
    ws2 = wb.create_sheet("Detailed Analysis")
    ws2.sheet_properties.tabColor = "70AD47"

    ws2.column_dimensions['A'].width = 28
    for col_letter in ['B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J']:
        ws2.column_dimensions[col_letter].width = 16
    ws2.column_dimensions['K'].width = 25
    ws2.column_dimensions['L'].width = 16

    r = 1

    # ── Section 1: Pro Forma Summary ──
    ws2.cell(row=r, column=1, value="PRO FORMA SUMMARY").font = Font(bold=True, size=14)
    r += 1

    cost_levels = [build_cost_psf - 25, build_cost_psf, build_cost_psf + 25, build_cost_psf + 50]
    _header_row(ws2, r, ["Line Item"] + [f"${cl}/sf" for cl in cost_levels] + ["", "Model Assumptions", "Value"])
    r += 1

    model_assumptions_right = [
        ("Units", units),
        ("Sq Ft / Unit", per_unit_sf),
        ("Total Sq Ft", build_sf),
        ("Land Cost", purchase_price),
        ("Exit $/sf", exit_psf),
        ("Build Duration (mo)", build_months),
        ("Hold Period (mo)", hold_months),
        ("Hard Contingency %", hard_contingency_pct / 100),
        ("Soft Cost %", soft_cost_pct / 100),
        ("LTV %", ltv / 100),
        ("Construction Rate %", interest_rate / 100),
        ("Perm Mortgage %", perm_mortgage_rate / 100),
    ]

    def _proforma_row_at_cost(bc):
        hc_l = bc * build_sf
        hcont_l = hc_l * (hard_contingency_pct / 100)
        sc_l = hc_l * (soft_cost_pct / 100)
        tdc_l = hc_l + hcont_l + sc_l + soft_contingency
        tpc_l = purchase_price + tdc_l
        la_l = tdc_l * (ltv / 100)
        ci_l = la_l * (interest_rate / 100) * (draw_factor / 100) * (build_months + delay_months) / 12
        lf_l = la_l * (loan_fee_pct / 100)
        carry_l = ci_l + lf_l
        rev_l = exit_psf * build_sf
        ec_l = rev_l * (exit_cost_pct / 100)
        total_l = tpc_l + carry_l + ec_l
        profit_l = rev_l - total_l + (net_rental_income if hold_months > 0 else 0)
        return {
            "Land": purchase_price,
            "Hard Cost": hc_l,
            f"Hard Cost Contingency ({hard_contingency_pct}%)": hcont_l,
            "Soft + Arch": sc_l + soft_contingency,
            "Soft Contingency": soft_contingency,
            "Carry (Interest + Fees)": carry_l,
            "Sales Cost": ec_l,
            "Total Cost": total_l,
            "Exit Value": rev_l,
            "Profit": profit_l,
        }

    proforma_rows = list(_proforma_row_at_cost(cost_levels[0]).keys())
    proforma_data = [_proforma_row_at_cost(cl) for cl in cost_levels]

    for i, row_label in enumerate(proforma_rows):
        vals = [row_label] + [proforma_data[j][row_label] for j in range(4)]
        fill_r = orange_fill if row_label == "Profit" else (light_blue_fill if row_label == "Total Cost" else None)
        fmts_r = [None] + [dollar_fmt] * 4

        # Add model assumptions on the right
        k_val = ""
        l_val = ""
        if i < len(model_assumptions_right):
            k_val, l_val = model_assumptions_right[i]
        vals += ["", k_val, l_val]
        k_fmt = None
        if isinstance(l_val, float) and l_val < 1:
            k_fmt = pct_fmt
        elif isinstance(l_val, (int, float)) and l_val >= 100:
            k_fmt = dollar_fmt
        fmts_r += [None, None, k_fmt]

        fills_r = [None] + [fill_r] * 4 + [None, None, yellow_fill if l_val != "" else None]
        _data_row(ws2, r, vals, fills=fills_r, fmts=fmts_r)
        r += 1

    # ── Section 2: Sales Cost Breakdown ──
    r += 1
    ws2.cell(row=r, column=1, value="SALES COST BREAKDOWN").font = Font(bold=True, size=14)
    r += 1
    _header_row(ws2, r, ["Category", "Rate", "Amount"])
    r += 1

    rev = exit_psf * build_sf
    sales_items = [
        ("Realtor / Agent", broker_fee_pct / 100, rev * broker_fee_pct / 100),
        ("Title + Closing", title_closing_pct / 100, rev * title_closing_pct / 100),
        ("Seller Concessions", seller_concessions_pct / 100, rev * seller_concessions_pct / 100),
        ("Total", exit_cost_pct / 100, exit_costs_user),
    ]
    for label, rate, amt in sales_items:
        fill_r = light_blue_fill if label == "Total" else None
        _data_row(ws2, r, [label, rate, amt],
                  fills=[fill_r, fill_r, fill_r],
                  fmts=[None, pct_fmt, dollar_fmt])
        r += 1

    # ── Section 3: Soft Cost Breakdown ──
    r += 1
    ws2.cell(row=r, column=1, value="SOFT COST BREAKDOWN").font = Font(bold=True, size=14)
    r += 1
    _header_row(ws2, r, ["Category", "Rate"] + [f"${cl}/sf" for cl in cost_levels])
    r += 1

    soft_categories = [
        ("Architecture & Design", arch_pct),
        ("Engineering (struct/MEP)", eng_pct),
        ("Permits & Impact Fees", permit_fee_pct),
        ("Surveys & Geotech", survey_pct),
        ("Builder's Risk Insurance", insurance_dev_pct),
        ("Other Soft Costs", other_soft_pct),
        ("Total Soft %", soft_cost_pct),
    ]
    for label, pct_val in soft_categories:
        hc_vals = [pct_val / 100] + [(cl * build_sf) * pct_val / 100 for cl in cost_levels]
        fill_r = light_blue_fill if "Total" in label else None
        _data_row(ws2, r, [label] + hc_vals,
                  fills=[fill_r, fill_r] + [fill_r] * 4,
                  fmts=[None, pct_fmt] + [dollar_fmt] * 4)
        r += 1

    # ── Section 4: Carrying Cost Breakdown ──
    r += 1
    ws2.cell(row=r, column=1, value="CARRYING COST BREAKDOWN").font = Font(bold=True, size=14)
    r += 1
    _header_row(ws2, r, ["Category", "Amount", "Notes"])
    r += 1

    carry_items = [
        ("Loan Interest (Construction)", construction_interest, f"{interest_rate}% × {draw_factor}% draw"),
        ("Loan Fees", loan_fees, f"{loan_fee_pct}% of loan"),
        ("Hold Interest (Perm Debt Service)", hold_interest, f"{hold_months} months"),
        ("Property Taxes (Hold)", prop_tax, f"{prop_tax_rate}%"),
        ("Insurance (Hold)", insurance, f"{hold_months} months"),
        ("Utilities / Misc (Hold)", misc, ""),
        ("Total Carry", construction_interest + loan_fees + hold_interest + prop_tax + insurance + misc, ""),
    ]
    for label, amt, note in carry_items:
        fill_r = light_blue_fill if "Total" in label else None
        _data_row(ws2, r, [label, amt, note],
                  fills=[fill_r, fill_r, fill_r],
                  fmts=[None, dollar_fmt, None])
        r += 1

    # ── Section 5: Cost vs Price Profit Matrix ──
    r += 1
    ws2.cell(row=r, column=1, value="COST vs PRICE PROFIT MATRIX").font = Font(bold=True, size=14)
    r += 1

    ws2.cell(row=r, column=1, value="Build Cost \\ Exit $/sf").font = bold_white
    ws2.cell(row=r, column=1).fill = header_fill
    for j, ep in enumerate(sweep_exit_prices, 2):
        cell = ws2.cell(row=r, column=j, value=f"${ep}/sf")
        cell.font = bold_white
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
    r += 1

    for bc in sweep_build_costs:
        ws2.cell(row=r, column=1, value=f"${bc}/sf").font = bold
        for j, ep in enumerate(sweep_exit_prices, 2):
            profit_val, _ = _compute_hold_profit(build_cost=bc, exit_price=ep, **common_args)
            cell = ws2.cell(row=r, column=j, value=profit_val)
            cell.number_format = dollar_fmt
            cell.fill = green_fill if profit_val >= 0 else PatternFill("solid", fgColor="FFC7CE")
            cell.border = thin_border
        r += 1

    # ── Section 6: Timeline Sensitivity ──
    r += 1
    ws2.cell(row=r, column=1, value="TIMELINE SENSITIVITY — Profit by Build Duration").font = Font(bold=True, size=14)
    r += 1

    durations = list(range(9, 16))
    ws2.cell(row=r, column=1, value="Duration (mo) \\ Build Cost").font = bold_white
    ws2.cell(row=r, column=1).fill = header_fill
    for j, bc in enumerate(cost_levels, 2):
        cell = ws2.cell(row=r, column=j, value=f"${bc}/sf")
        cell.font = bold_white
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
    r += 1

    for dur in durations:
        ws2.cell(row=r, column=1, value=f"{dur} months").font = bold
        for j, bc in enumerate(cost_levels, 2):
            profit_val, _ = _compute_hold_profit(
                build_cost=bc, exit_price=exit_psf,
                purchase_price=purchase_price, build_sf=build_sf,
                exit_cost_pct=exit_cost_pct,
                hard_contingency_pct=hard_contingency_pct,
                soft_cost_pct=soft_cost_pct, soft_contingency=soft_contingency,
                ltv=ltv, interest_rate=interest_rate, draw_factor=draw_factor,
                loan_fee_pct=loan_fee_pct,
                build_months=dur, hold_months=hold_months, delay_months=delay_months,
                perm_mortgage_rate=perm_mortgage_rate, amortization_years=amortization_years,
                rent_per_unit=rent_per_unit, units=units, vacancy_pct=vacancy_pct,
                mgmt_fee_pct=mgmt_fee_pct, prop_tax_rate=prop_tax_rate,
                taxable_value_psf=taxable_value_psf, insurance_monthly=insurance_monthly,
                repairs_per_unit=repairs_per_unit, common_utilities=common_utilities,
                leasing_reserve=leasing_reserve,
            )
            cell = ws2.cell(row=r, column=j, value=profit_val)
            cell.number_format = dollar_fmt
            cell.fill = green_fill if profit_val >= 0 else PatternFill("solid", fgColor="FFC7CE")
            cell.border = thin_border
        r += 1

    # Save to bytes
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════
#  STREAMLIT UI — Combined Due Diligence + Financial Modeling
# ══════════════════════════════════════════════════════════════

st.title("🏡 Austin Deal Analyzer PRO")
st.caption("Real market data + financial modeling → **Should you buy or not?**")

# ── Sidebar: Input Form ──
with st.sidebar:
    st.header("📥 Deal Inputs")

    with st.form("deal_form"):
        st.subheader("🏠 Property")
        address = st.text_input("Property Address", value=st.session_state.get('address', ''),
                               placeholder="e.g., 2613 Nottingham Ln",
                               help="Street address of the property you're analyzing")
        zip_code = st.text_input("ZIP Code", value=st.session_state.get('zip_code', ''),
                                 placeholder="e.g., 78704",
                                 help="Used to pull comps, permits, and zoning data")

        st.divider()
        st.subheader("💵 Deal Numbers")
        purchase_price = st.number_input("Purchase Price ($)", min_value=0, value=450000, step=25000,
                                         help="Land acquisition cost or total purchase price")
        build_sf = st.number_input("Total Build Size (sf)", min_value=0, value=3000, step=500,
                                   help="Total finished square footage across all units")
        units = st.number_input("Number of Units", min_value=1, value=2, step=1,
                                help="Number of residential units (e.g., 2 for a duplex)")
        build_cost_psf = st.number_input("Build Cost ($/sf)", min_value=0, value=250, step=25,
                                         help="Hard construction cost per square foot (labor + materials)")
        exit_psf = st.number_input("Exit Price ($/sf)", min_value=0, value=575, step=10,
                                   help="Your target sale price per square foot")

        st.divider()
        st.subheader("💵 Cost Details")
        hard_contingency_pct = st.slider("Hard Cost Contingency (%)", 0, 15, 6,
                                         help="Buffer for unexpected construction cost overruns (typically 5-10%)")
        split_soft = st.toggle("Split Soft Cost Categories", value=False,
                               help="Break down soft costs into individual line items instead of one percentage")
        if split_soft:
            arch_pct = st.number_input("Architecture & Design (%)", min_value=0.0, max_value=10.0, value=3.0, step=0.5,
                                       help="Architect fees — as % of hard cost")
            eng_pct = st.number_input("Engineering (structural/MEP) (%)", min_value=0.0, max_value=10.0, value=2.0, step=0.5,
                                      help="Structural, mechanical, electrical, plumbing engineering — as % of hard cost")
            permit_fee_pct = st.number_input("Permits & Impact Fees (%)", min_value=0.0, max_value=10.0, value=2.5, step=0.5,
                                             help="City permits, impact fees, utility connections — as % of hard cost")
            survey_pct = st.number_input("Surveys & Geotech (%)", min_value=0.0, max_value=5.0, value=1.0, step=0.5,
                                         help="Land survey, soil testing, environmental — as % of hard cost")
            insurance_dev_pct = st.number_input("Builder's Risk Insurance (%)", min_value=0.0, max_value=5.0, value=1.0, step=0.5,
                                                help="Builder's risk / liability during construction — as % of hard cost")
            other_soft_pct = st.number_input("Other Soft Costs (%)", min_value=0.0, max_value=10.0, value=0.9, step=0.1,
                                             help="Legal, accounting, misc — as % of hard cost")
            soft_cost_pct = arch_pct + eng_pct + permit_fee_pct + survey_pct + insurance_dev_pct + other_soft_pct
            st.caption(f"**Total Soft: {soft_cost_pct:.1f}%**")
        else:
            soft_cost_pct = st.number_input("Soft Costs (arch/eng/permits) (%)", min_value=0.0, max_value=30.0, value=10.4, step=0.5,
                                            help="Architecture, engineering, permits, surveys — as % of hard cost")
            arch_pct = eng_pct = permit_fee_pct = survey_pct = insurance_dev_pct = other_soft_pct = 0.0
        soft_contingency = st.number_input("Soft Contingency ($)", min_value=0, value=20000, step=5000,
                                           help="Fixed buffer for unexpected soft cost items")

        st.divider()
        st.subheader("💰 Construction Financing")
        ltv = st.slider("Loan to Cost (%)", 0, 100, 100,
                        help="% of non-land development cost funded by debt")
        interest_rate = st.slider("Construction Interest Rate (%)", 3.0, 14.0, 8.0, step=0.5,
                                  help="Annual interest rate on construction loan")
        draw_factor = st.slider("Draw Factor (%)", 40, 80, 60,
                                help="Avg % of loan funded during construction")
        loan_fee_pct = st.slider("Construction Loan Fees (%)", 0.0, 3.0, 1.0, step=0.25,
                                 help="Origination / lender fees on construction debt")

        st.divider()
        st.subheader("📅 Timeline")
        build_months = st.slider("Build Duration (months)", 6, 24, 12,
                                help="Estimated construction timeline from permit to CO")
        hold_months = st.slider("Hold Period After Build (months)", 0, 36, 24,
                                help="0 = flip immediately, 24 = rent then sell")
        delay_months = st.slider("Expected Delays (months)", 0, 12, 0,
                                help="Buffer for permitting delays, weather, supply issues")

        st.divider()
        st.subheader("🏘️ Rental (Hold Strategy)")
        rent_per_unit = st.number_input("Monthly Rent / Unit ($)", min_value=0, value=3950, step=100,
                                      help="Expected monthly rent per unit after lease-up")
        vacancy_pct = st.slider("Vacancy / Credit Loss (%)", 0, 15, 5,
                                help="% of gross rent lost to vacancy and bad debt")
        mgmt_fee_pct = st.slider("Management Fee (%)", 0, 15, 7,
                                 help="Property management fee as % of effective rent")
        perm_mortgage_rate = st.slider("Permanent Mortgage Rate (%)", 3.0, 12.0, 7.0, step=0.25,
                                       help="Rate after construction loan converts to permanent")
        amortization_years = st.number_input("Amortization (years)", min_value=15, max_value=30, value=30, step=5,
                                            help="Loan payoff schedule length (longer = lower monthly payment)")
        taxable_value_psf = st.number_input("Taxable Value ($/sf)", min_value=0, value=550, step=25,
                                            help="Assessed value for property tax during hold")
        prop_tax_rate = st.slider("Property Tax Rate (%)", 1.0, 4.0, 2.0, step=0.1,
                                  help="Annual property tax rate (Austin is typically ~2%)")
        insurance_monthly = st.number_input("Landlord Insurance ($/mo)", min_value=0, value=375, step=25,
                                            help="Monthly hazard + liability insurance premium")
        repairs_per_unit = st.number_input("Repairs Reserve ($/unit/mo)", min_value=0, value=150, step=25,
                                           help="Monthly reserve per unit for maintenance and repairs")
        common_utilities = st.number_input("Common Utilities / Misc ($/mo)", min_value=0, value=250, step=25,
                                           help="Owner-paid utilities, landscaping, pest control, etc.")
        leasing_reserve = st.number_input("Leasing / Turnover Reserve ($/mo)", min_value=0, value=250, step=25,
                                          help="Reserve for tenant turnover, marketing, and lease-up costs")

        st.divider()
        st.subheader("📉 Market Risk")
        price_decline = st.slider("Annual Price Change (%)", -15, 10, 0,
                                  help="Expected annual change in market prices (negative = decline)")

        st.divider()
        st.subheader("🏗️ Construction Carry Costs")
        const_tax_rate = st.slider("Construction Property Tax (%)", 0.0, 4.0, 2.0, step=0.1,
                                    help="Annual property tax rate during construction period")
        const_insurance_annual = st.number_input("Construction Insurance ($/yr)", min_value=0, value=6750, step=250,
                                                  help="Builder's risk + liability insurance per year during construction")
        const_utilities = st.number_input("Construction Utilities ($/mo)", min_value=0, value=450, step=50,
                                           help="Water, electric, temp power during construction")
        const_misc = st.number_input("Construction Misc ($/mo)", min_value=0, value=325, step=25,
                                      help="Dumpster, portable toilet, misc during construction")
        carry_buffer_pct = st.slider("Carry Cost Buffer (%)", 0, 25, 12, step=1,
                                      help="Buffer on top of all carry costs for unexpected overruns")

        st.divider()
        st.subheader("💸 Sale / Exit Costs")
        broker_fee_pct = st.slider("Broker / Agent Fee (%)", 0.0, 6.0, 3.0, step=0.5,
                                   help="Listing + buyer agent commission")
        title_closing_pct = st.slider("Title + Closing Costs (%)", 0.0, 3.0, 1.5, step=0.25,
                                      help="Title insurance, escrow, recording fees")
        seller_concessions_pct = st.slider("Seller Concessions (%)", 0.0, 3.0, 0.5, step=0.25,
                                           help="Buyer credits, repairs, warranty")
        exit_cost_pct = broker_fee_pct + title_closing_pct + seller_concessions_pct
        sale_hold_months = st.number_input('Sale Hold Period (months)', min_value=0.0, max_value=6.0, value=1.5, step=0.5,
                                            help='Months property sits on market before closing')
        staging_base = st.number_input('Staging Base ($)', min_value=0, value=1500, step=500,
                                        help='Base staging cost (fixed)')
        staging_per_unit = st.number_input('Staging Per Unit ($)', min_value=0, value=3500, step=500,
                                            help='Additional staging cost per unit')
        marketing_base = st.number_input('Marketing Base ($)', min_value=0, value=2000, step=500,
                                          help='Photography, signage, MLS listing fees')
        marketing_per_unit = st.number_input('Marketing Per Unit ($)', min_value=0, value=1500, step=500,
                                              help='Additional marketing cost per unit')
        warranty_per_unit = st.number_input('Warranty Per Unit ($)', min_value=0, value=1750, step=250,
                                             help='Home warranty cost per unit')

        submitted = st.form_submit_button("🔍 Analyze — Should I Buy?", use_container_width=True, type="primary")


# ── Main Content ──
show_analysis = submitted and address and zip_code
result = None

if submitted and address and zip_code:
    st.session_state['address'] = address
    st.session_state['zip_code'] = zip_code
    st.session_state['analysis_done'] = True
    street_name = extract_street_name(address)

    # ══════════════════════════════════════════════
    # STEP 1: Pull real market data
    # ══════════════════════════════════════════════
    with st.spinner(f"Pulling real market data for {address}, {zip_code}..."):
        result = run_analysis(address, zip_code, street_name)
    st.session_state['result'] = result
elif st.session_state.get('analysis_done') and not submitted:
    # Rerun triggered by button click (e.g. AI Analysis) — restore cached result
    show_analysis = True
    result = st.session_state.get('result')
    street_name = extract_street_name(address)

if show_analysis and result is not None:

    # ══════════════════════════════════════════════
    # STEP 2: Financial calculations
    # ══════════════════════════════════════════════
    hard_cost = build_cost_psf * build_sf
    hard_contingency = hard_cost * (hard_contingency_pct / 100)
    soft_costs = hard_cost * (soft_cost_pct / 100)
    total_dev_cost = hard_cost + hard_contingency + soft_costs + soft_contingency
    total_project_cost = purchase_price + total_dev_cost

    # Construction financing
    loan_amount = total_dev_cost * (ltv / 100)  # LTC on development costs (not land)
    equity = total_project_cost - loan_amount
    total_months = build_months + hold_months + delay_months
    timeline_years = total_months / 12

    # Construction interest (draw factor applies during build only)
    construction_interest = loan_amount * (interest_rate / 100) * (draw_factor / 100) * (build_months + delay_months) / 12
    loan_fees = loan_amount * (loan_fee_pct / 100)

    # Construction carry costs (during build, matching Karen Ave Excel)
    carry_months = build_months + delay_months
    carry_loan_interest = (purchase_price + total_dev_cost) * (ltv / 100) * (interest_rate / 100) * (draw_factor / 100) * carry_months / 12
    carry_taxes = (purchase_price + 0.5 * total_dev_cost) * (const_tax_rate / 100) * carry_months / 12
    carry_insurance = const_insurance_annual * carry_months / 12
    carry_utilities = const_utilities * carry_months
    carry_misc = const_misc * carry_months
    carry_subtotal = carry_loan_interest + carry_taxes + carry_insurance + carry_utilities + carry_misc
    carry_buffer = carry_subtotal * (carry_buffer_pct / 100)
    total_carry = carry_subtotal + carry_buffer

    # Sales costs (matching Karen Ave Excel)
    user_revenue_est = exit_psf * build_sf
    staging_cost = staging_base + staging_per_unit * units
    marketing_cost = marketing_base + marketing_per_unit * units
    warranty_cost = warranty_per_unit * units
    variable_sales = user_revenue_est * (exit_cost_pct / 100)
    holding_during_sale = total_carry / max(carry_months, 1) * sale_hold_months
    total_sales_cost = variable_sales + staging_cost + marketing_cost + warranty_cost + holding_during_sale

    # Holding costs during rental period
    if hold_months > 0:
        # Permanent mortgage debt service (amortized PMT on construction debt)
        monthly_perm_rate = perm_mortgage_rate / 100 / 12
        n_payments = amortization_years * 12
        if monthly_perm_rate > 0:
            monthly_debt_service = loan_amount * (monthly_perm_rate * (1 + monthly_perm_rate) ** n_payments) / ((1 + monthly_perm_rate) ** n_payments - 1)
        else:
            monthly_debt_service = loan_amount / n_payments
        hold_interest = monthly_debt_service * hold_months

        # Loan balance after hold (amortized)
        if monthly_perm_rate > 0:
            loan_balance_after_hold = loan_amount * (1 + monthly_perm_rate) ** hold_months - monthly_debt_service * ((1 + monthly_perm_rate) ** hold_months - 1) / monthly_perm_rate
        else:
            loan_balance_after_hold = loan_amount - (monthly_debt_service * hold_months)

        # Rental income during hold
        gross_rent = rent_per_unit * units * hold_months
        effective_rent = gross_rent * (1 - vacancy_pct / 100)
        mgmt_cost = effective_rent * (mgmt_fee_pct / 100)
        # Property tax during hold (based on taxable value)
        prop_tax = build_sf * taxable_value_psf * (prop_tax_rate / 100) * hold_months / 12
        # Operating expenses during hold
        insurance = insurance_monthly * hold_months
        repairs = repairs_per_unit * units * hold_months
        misc = common_utilities * hold_months
        leasing = leasing_reserve * hold_months
        total_hold_expenses = hold_interest + mgmt_cost + prop_tax + insurance + repairs + misc + leasing

        # Monthly NOI (before debt service)
        monthly_noi = (effective_rent - mgmt_cost - prop_tax - insurance - repairs - misc - leasing) / max(hold_months, 1)
        monthly_cf_after_debt = monthly_noi - monthly_debt_service

        net_rental_income = effective_rent - total_hold_expenses
    else:
        hold_interest = 0
        gross_rent = 0
        effective_rent = 0
        net_rental_income = 0
        total_hold_expenses = 0
        loan_balance_after_hold = loan_amount
        monthly_debt_service = 0
        monthly_noi = 0
        monthly_cf_after_debt = 0
        prop_tax = 0
        insurance = 0
        repairs = 0
        misc = 0
        leasing = 0
        mgmt_cost = 0

    total_interest = construction_interest + hold_interest

    # Use REAL market median if available, otherwise use user's exit assumption
    per_unit_sf = build_sf / max(units, 1)
    # Filter comps to similar per-unit size (±30%)
    similar_comps = [c for c in result.redfin_comps
                     if c.get("sqft", 0) > 0 and abs(c["sqft"] - per_unit_sf) / per_unit_sf <= 0.30]
    if similar_comps:
        sim_psf = sorted([c["psf"] for c in similar_comps if c.get("psf", 0) > 0])
        if sim_psf:
            sim_mid = len(sim_psf) // 2
            result.similar_stats = {
                "median_psf": sim_psf[sim_mid],
                "avg_psf": round(sum(sim_psf) / len(sim_psf)),
                "min_psf": min(sim_psf),
                "max_psf": max(sim_psf),
                "count": len(sim_psf),
                "size_range": f"{int(per_unit_sf * 0.7):,}–{int(per_unit_sf * 1.3):,} sf",
            }
    similar_median = getattr(result, 'similar_stats', {}).get('median_psf', 0)
    all_median = result.market_stats.get("median_psf", 0)
    # Prefer similar-size median for market comparison
    median_psf = similar_median if similar_median > 0 else all_median
    # When Census fallback (no real comps), use user's exit price for market calcs
    # Census median includes all homes (old, small) — not valid for new construction pricing
    is_census_fallback = result.market_stats.get('source', '').startswith('Census') if result.market_stats else False
    if is_census_fallback:
        market_exit = exit_psf
    else:
        market_exit = median_psf if median_psf > 0 else exit_psf

    # Adjusted exit with market trend
    price_change_rate = price_decline / 100
    adjusted_exit = market_exit * ((1 + price_change_rate) ** timeline_years)
    adjusted_revenue = adjusted_exit * build_sf
    user_revenue = exit_psf * build_sf

    # Exit costs (realtor, title, closing)
    exit_costs_user = user_revenue * (exit_cost_pct / 100)
    exit_costs_market = adjusted_revenue * (exit_cost_pct / 100)

    # Profit calculation — Equity/Debt model (matches Excel)
    if hold_months > 0:
        # Total equity invested = Land + Construction Interest + Loan Fees + Additional Equity for negative CF
        cumulative_cf = monthly_cf_after_debt * hold_months
        additional_equity_needed = max(0, -cumulative_cf)
        total_equity_invested = purchase_price + construction_interest + loan_fees + additional_equity_needed

        # At sale: pay off loan balance, keep net proceeds + any positive rental CF
        net_sale_before_debt = user_revenue - exit_costs_user
        net_sale_after_debt = net_sale_before_debt - loan_balance_after_hold
        positive_rental_cf = max(0, cumulative_cf)
        total_cash_returned = net_sale_after_debt + positive_rental_cf

        user_profit = total_cash_returned - total_equity_invested
        equity_multiple = total_cash_returned / total_equity_invested if total_equity_invested > 0 else 0

        # Market-based profit
        net_market_sale = adjusted_revenue - exit_costs_market
        market_net_after_debt = net_market_sale - loan_balance_after_hold
        market_cash_returned = market_net_after_debt + positive_rental_cf
        market_profit = market_cash_returned - total_equity_invested
    else:
        # Simple flip model
        total_cost = total_project_cost + construction_interest + loan_fees + exit_costs_user
        total_equity_invested = total_cost
        user_profit = user_revenue - total_cost
        market_profit = adjusted_revenue - (total_project_cost + construction_interest + loan_fees + exit_costs_market)
        equity_multiple = user_revenue / total_cost if total_cost > 0 else 0
        cumulative_cf = 0
        additional_equity_needed = 0
        total_cash_returned = user_revenue - exit_costs_user

    total_cost = total_equity_invested  # for break-even calc
    total_interest = construction_interest + (hold_interest if hold_months > 0 else 0)

    annual_rent = rent_per_unit * 12 * units
    cash_yield = (annual_rent / total_equity_invested * 100) if total_equity_invested > 0 else 0
    annualized_return = ((total_cash_returned / total_equity_invested) ** (1 / max(timeline_years, 0.5)) - 1) if total_equity_invested > 0 else 0
    breakeven_psf = total_equity_invested / build_sf if build_sf > 0 else 0

    # ══════════════════════════════════════════════
    # STEP 3: Risk scoring (0-100, higher = more risk)
    # ══════════════════════════════════════════════
    risk_score = 0
    risk_flags = []

    active_permits = [p for p in result.street_permits if p.status == 'Active']

    # Market data risks
    comp_count = result.market_stats.get('count', 0) if result.market_stats else 0
    if median_psf > 0 and not is_census_fallback:
        exit_gap = ((exit_psf - median_psf) / median_psf) * 100
        if exit_gap > 20:
            risk_score += 30
            risk_flags.append(("🔴", f"Exit ${exit_psf}/sf is **{exit_gap:.0f}% above** market median ${median_psf}/sf — UNREALISTIC"))
        elif exit_gap > 10:
            risk_score += 15
            risk_flags.append(("🟡", f"Exit ${exit_psf}/sf is **{exit_gap:.0f}% above** market median ${median_psf}/sf — AGGRESSIVE"))
        elif exit_gap > 0:
            risk_score += 5
            risk_flags.append(("🟢", f"Exit ${exit_psf}/sf is **{exit_gap:.0f}% above** market median ${median_psf}/sf — Reasonable"))
        else:
            risk_flags.append(("🟢", f"Exit ${exit_psf}/sf is **at or below** market median ${median_psf}/sf — Conservative"))
    elif is_census_fallback:
        risk_score += 10
        risk_flags.append(("🟡", f"**No Redfin comps** — Census estimate ${median_psf}/sf includes all homes (not just new construction). Verify exit price with actual comps."))
    else:
        risk_score += 30
        risk_flags.append(("🔴", "**No market comp data** — cannot validate exit assumptions. Redfin may be blocking this server."))

    # Competition risks
    if len(active_permits) >= 5:
        risk_score += 20
        risk_flags.append(("🔴", f"**{len(active_permits)} competing units** under active construction on {street_name} St"))
    elif len(active_permits) >= 3:
        risk_score += 10
        risk_flags.append(("🟡", f"{len(active_permits)} competing units under construction on {street_name} St"))
    elif len(active_permits) > 0:
        risk_flags.append(("🟢", f"{len(active_permits)} unit(s) under construction — manageable competition"))

    # Profitability risks
    if market_profit < 0:
        risk_score += 25
        risk_flags.append(("🔴", f"**Negative profit** at market exit (${adjusted_exit:.0f}/sf) — LOSS of ${abs(market_profit):,.0f}"))
    elif market_profit < 100000:
        risk_score += 15
        risk_flags.append(("🟡", f"Thin profit margin (${market_profit:,.0f}) — vulnerable to overruns"))

    if delay_months > 3:
        risk_score += 10
        risk_flags.append(("🟡", f"Delays adding ${delay_cost:,.0f} in holding costs"))

    if breakeven_psf > median_psf and median_psf > 0 and not is_census_fallback:
        risk_score += 20
        risk_flags.append(("🔴", f"Break-even (${breakeven_psf:.0f}/sf) is **above** market median (${median_psf}/sf)"))

    # ══════════════════════════════════════════════
    # STEP 4: THE VERDICT
    # ══════════════════════════════════════════════
    st.markdown("---")

    if risk_score >= 50 or market_profit < 0:
        verdict = "DON'T BUY"
        verdict_color = "red"
        verdict_emoji = "❌"
        verdict_detail = "Too many risk factors. This deal doesn't pencil at current market rates."
    elif risk_score >= 25 or market_profit < 100000 or median_psf == 0:
        verdict = "CAUTION"
        verdict_color = "orange"
        verdict_emoji = "⚠️"
        if median_psf == 0:
            verdict_detail = "No market data to validate assumptions. Cannot confirm this is a good deal — research comps manually."
        else:
            verdict_detail = "Deal is marginal. Only proceed if you can negotiate better terms."
    else:
        verdict = "BUY"
        verdict_color = "green"
        verdict_emoji = "✅"
        verdict_detail = "Deal looks solid based on market data and financials."

    # Check listing status — override verdict if not available
    listing = result.listing_status
    listing_status = listing.get('status', 'Unknown')
    listing_url = listing.get('url', '')

    if listing_status == 'Pending':
        verdict_detail += " ⚠️ BUT this property is PENDING (under contract) — may not be available."
        verdict_emoji = "🔒" if verdict == "BUY" else verdict_emoji
        if verdict == "BUY":
            verdict = "UNDER CONTRACT"
            verdict_color = "orange"
    elif listing_status == 'Sold':
        verdict_detail += " 🏠 This property has already SOLD."
        verdict_emoji = "🔒" if verdict == "BUY" else verdict_emoji
        if verdict == "BUY":
            verdict = "ALREADY SOLD"
            verdict_color = "orange"

    # Big verdict banner
    st.markdown(f"""
    <div style="background-color: {'#ff4b4b' if verdict == "DON'T BUY" else '#ffa726' if verdict in ('CAUTION', 'UNDER CONTRACT', 'ALREADY SOLD') else '#4caf50'};
                padding: 30px; border-radius: 15px; text-align: center; margin: 10px 0 20px 0;">
        <h1 style="color: white; margin: 0; font-size: 48px;">{verdict_emoji} {verdict}</h1>
        <p style="color: white; margin: 10px 0 0 0; font-size: 18px;">{verdict_detail}</p>
    </div>
    """, unsafe_allow_html=True)

    # Key numbers
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("All-In Cost", f"${total_cost:,.0f}")
    c2.metric("Market Exit Revenue", f"${adjusted_revenue:,.0f}")
    c3.metric("Profit (Market)", f"${market_profit:,.0f}")
    c4.metric("Break-Even", f"${breakeven_psf:.0f}/sf")
    c5.metric("Risk Score", f"{risk_score}/100")

    if median_psf > 0:
        if is_census_fallback:
            st.info(f"📊 **Market Data:** No Redfin comps available. Census estimate: **${median_psf}/sf** "
                    f"(includes all homes, not just new construction). "
                    f"Your exit: ${exit_psf}/sf. Using your exit price for profit calculations.")
        else:
            median_source = f"similar-size ({int(per_unit_sf):,} sf ±30%)" if similar_median > 0 else "all comps"
            sim_count = getattr(result, 'similar_stats', {}).get('count', 0)
            all_count = result.market_stats.get('count', 0)
            st.info(f"📊 **Market Data:** Median for {median_source} is **${median_psf}/sf** "
                    f"({sim_count if similar_median > 0 else all_count} comps). "
                    f"All comps median: ${all_median}/sf ({all_count}). "
                    f"Your exit: ${exit_psf}/sf.")

    st.markdown("---")

    # ══════════════════════════════════════════════
    # TABS
    # ══════════════════════════════════════════════
    tab_verdict, tab_scenarios, tab_comps, tab_active, tab_rental, tab_permits, tab_plot, tab_ai, tab_download = st.tabs([
        "🚦 Risk Analysis", "📈 Scenarios & Sensitivity", "📊 Sold Comps", "🏠 Active Listings", "💰 Rental Comps", "🏗️ Permits", "📋 Plot Info", "🤖 AI Analysis", "📄 Download"
    ])

    # ── Risk Analysis Tab ──
    with tab_verdict:
        st.subheader("Risk Factors")
        for emoji, msg in risk_flags:
            st.markdown(f"{emoji} {msg}")

        st.markdown("---")
        st.subheader("Deal Structure")

        c1, c2 = st.columns(2)
        with c1:
            st.markdown("**Development & Financing**")
            soft_cost_line = ""
            if split_soft:
                soft_cost_line = f"""| — Architecture & Design ({arch_pct}%) | ${hard_cost * arch_pct / 100:,.0f} |
            | — Engineering ({eng_pct}%) | ${hard_cost * eng_pct / 100:,.0f} |
            | — Permits & Impact Fees ({permit_fee_pct}%) | ${hard_cost * permit_fee_pct / 100:,.0f} |
            | — Surveys & Geotech ({survey_pct}%) | ${hard_cost * survey_pct / 100:,.0f} |
            | — Builder's Risk Insurance ({insurance_dev_pct}%) | ${hard_cost * insurance_dev_pct / 100:,.0f} |
            | — Other Soft ({other_soft_pct}%) | ${hard_cost * other_soft_pct / 100:,.0f} |
            | **Soft Costs Total ({soft_cost_pct:.1f}%)** | **${soft_costs:,.0f}** |"""
            else:
                soft_cost_line = f"| Soft Costs ({soft_cost_pct}%) | ${soft_costs:,.0f} |"
            st.markdown(f"""
            | Item | Amount |
            |------|--------|
            | Land / Purchase | ${purchase_price:,.0f} |
            | Hard Cost ({build_sf:,} sf × ${build_cost_psf}/sf) | ${hard_cost:,.0f} |
            | Hard Contingency ({hard_contingency_pct}%) | ${hard_contingency:,.0f} |
            {soft_cost_line}
            | Soft Contingency | ${soft_contingency:,.0f} |
            | **Non-Land Dev Cost** | **${total_dev_cost:,.0f}** |
            | Construction Debt ({ltv}% LTV) | ${loan_amount:,.0f} |
            | Construction Interest ({interest_rate}%) | ${construction_interest:,.0f} |
            | Construction Loan Fees ({loan_fee_pct}%) | ${loan_fees:,.0f} |
            """)

        with c2:
            if hold_months > 0:
                st.markdown("**Equity Investment & Returns**")
                st.markdown(f"""
                | Item | Amount |
                |------|--------|
                | Land Equity | ${purchase_price:,.0f} |
                | Construction Interest | ${construction_interest:,.0f} |
                | Loan Fees | ${loan_fees:,.0f} |
                | Additional Equity (negative CF) | ${additional_equity_needed:,.0f} |
                | **Total Equity Invested** | **${total_equity_invested:,.0f}** |
                | | |
                | Sale Price (${exit_psf}/sf) | ${user_revenue:,.0f} |
                | Less Sale Costs ({exit_cost_pct:.1f}%) | -${exit_costs_user:,.0f} |
                | — Broker ({broker_fee_pct}%) | -${user_revenue * broker_fee_pct / 100:,.0f} |
                | — Title/Closing ({title_closing_pct}%) | -${user_revenue * title_closing_pct / 100:,.0f} |
                | — Concessions ({seller_concessions_pct}%) | -${user_revenue * seller_concessions_pct / 100:,.0f} |
                | Less Loan Payoff | -${loan_balance_after_hold:,.0f} |
                | Plus Positive Rental CF | +${max(0, cumulative_cf):,.0f} |
                | **Cash Returned** | **${total_cash_returned:,.0f}** |
                | **Profit** | **${user_profit:,.0f}** |
                | **Equity Multiple** | **{equity_multiple:.2f}x** |
                """)
            else:
                st.markdown("**Returns (Flip)**")
                st.markdown(f"""
                | Scenario | Revenue | Profit |
                |----------|---------|--------|
                | Your Exit (${exit_psf}/sf) | ${user_revenue:,.0f} | ${user_profit:,.0f} |
                | Market Exit (${adjusted_exit:.0f}/sf) | ${adjusted_revenue:,.0f} | ${market_profit:,.0f} |
                | Sale Costs ({exit_cost_pct:.1f}%) | -${exit_costs_user:,.0f} | |
                | — Broker ({broker_fee_pct}%) | -${user_revenue * broker_fee_pct / 100:,.0f} | |
                | — Title/Closing ({title_closing_pct}%) | -${user_revenue * title_closing_pct / 100:,.0f} | |
                | — Concessions ({seller_concessions_pct}%) | -${user_revenue * seller_concessions_pct / 100:,.0f} | |
                """)

        st.markdown("---")
        st.subheader("Hold Strategy (Rental)")
        if hold_months > 0:
            rc1, rc2, rc3, rc4, rc5 = st.columns(5)
            rc1.metric("Gross Rent", f"${gross_rent:,.0f}", delta=f"{hold_months} months")
            rc2.metric("Net Rental Income", f"${net_rental_income:,.0f}")
            rc3.metric("Monthly Cash Flow", f"${monthly_cf_after_debt:,.0f}")
            rc4.metric("Monthly Debt Service", f"${monthly_debt_service:,.0f}")
            rc5.metric("Loan Balance After Hold", f"${loan_balance_after_hold:,.0f}")

            st.markdown(f"""
            **Rental NOI Detail ({hold_months} months):**
            | Item | Monthly | Total ({hold_months} mo) |
            |------|---------|---------|
            | Gross Rent ({units} × ${rent_per_unit:,}/mo) | ${rent_per_unit * units:,.0f} | ${gross_rent:,.0f} |
            | Less Vacancy ({vacancy_pct}%) | -${rent_per_unit * units * vacancy_pct / 100:,.0f} | -${gross_rent * vacancy_pct / 100:,.0f} |
            | **Effective Gross Income** | **${effective_rent / hold_months:,.0f}** | **${effective_rent:,.0f}** |
            | Less Mgmt Fee ({mgmt_fee_pct}%) | -${mgmt_cost / hold_months:,.0f} | -${mgmt_cost:,.0f} |
            | Less Property Tax ({prop_tax_rate}%) | -${prop_tax / hold_months:,.0f} | -${prop_tax:,.0f} |
            | Less Insurance | -${insurance_monthly:,.0f} | -${insurance:,.0f} |
            | Less Repairs ({units} × ${repairs_per_unit}/mo) | -${repairs_per_unit * units:,.0f} | -${repairs:,.0f} |
            | Less Utilities/Misc | -${common_utilities:,.0f} | -${misc:,.0f} |
            | Less Leasing/Turnover | -${leasing_reserve:,.0f} | -${leasing:,.0f} |
            | **NOI (before debt)** | **${monthly_noi:,.0f}** | **${monthly_noi * hold_months:,.0f}** |
            | Less Debt Service ({perm_mortgage_rate}%, {amortization_years}yr) | -${monthly_debt_service:,.0f} | -${monthly_debt_service * hold_months:,.0f} |
            | **Cash Flow After Debt** | **${monthly_cf_after_debt:,.0f}** | **${monthly_cf_after_debt * hold_months:,.0f}** |
            """)

            if monthly_cf_after_debt < 0:
                st.warning(f"⚠️ **Negative cash flow** of ${monthly_cf_after_debt:,.0f}/mo during hold period. "
                          f"You'll need ${abs(monthly_cf_after_debt * hold_months):,.0f} additional equity to carry this property.")
        else:
            rc1, rc2, rc3 = st.columns(3)
            rc1.metric("Annual Rent (if held)", f"${annual_rent:,.0f}")
            rc2.metric("Cash-on-Cash Yield", f"{cash_yield:.1f}%")
            rc3.metric("Equity Invested", f"${equity:,.0f}")
            st.info("Hold period is 0 — this is a flip strategy. Set hold months > 0 to see rental analysis.")

        st.caption("⚠ TCAD data (appraisals, deeds, foreclosure checks) requires running the CLI tool locally for complete risk assessment.")

    # ── Scenarios & Sensitivity Tab ──
    with tab_scenarios:
        # ── Pro Forma Summary (matching Karen Ave Excel Page 1) ──
        st.subheader(f"📋 Pro Forma Summary (${exit_psf}/sf Exit)")
        st.caption("Full cost breakdown across multiple build cost scenarios")
        pf_cost_levels = sorted(set([max(100, build_cost_psf - 50), max(125, build_cost_psf - 25), build_cost_psf, build_cost_psf + 25, build_cost_psf + 50]))

        pf_data = []
        for row_label in ["Land", "Hard Cost", "Hard Contingency", "Soft + Arch", "Soft Contingency", "Carry", "Sales Cost", "Exit Value", "Total Cost", "**Profit**"]:
            row = {"Category": row_label}
            for bc in pf_cost_levels:
                hc = bc * build_sf
                hc_cont = hc * (hard_contingency_pct / 100)
                sc = hc * (soft_cost_pct / 100)
                dev = hc + hc_cont + sc + soft_contingency
                loan = dev * (ltv / 100)
                # Carry
                ci = loan * (interest_rate / 100) * (draw_factor / 100) * carry_months / 12
                ct = (purchase_price + 0.5 * dev) * (const_tax_rate / 100) * carry_months / 12
                cins = const_insurance_annual * carry_months / 12
                cutil = const_utilities * carry_months
                cmisc = const_misc * carry_months
                csub = ci + ct + cins + cutil + cmisc
                cbuf = csub * (carry_buffer_pct / 100)
                tcarry = csub + cbuf
                # Sales
                rev = exit_psf * build_sf
                ec = rev * (exit_cost_pct / 100)
                stg = staging_base + staging_per_unit * units
                mkt = marketing_base + marketing_per_unit * units
                war = warranty_per_unit * units
                hds = tcarry / max(carry_months, 1) * sale_hold_months
                tsc = ec + stg + mkt + war + hds
                lf = loan * (loan_fee_pct / 100)
                total = purchase_price + hc + hc_cont + sc + soft_contingency + tcarry + tsc + lf
                profit = rev - total

                val = {"Land": purchase_price, "Hard Cost": hc, "Hard Contingency": hc_cont,
                       "Soft + Arch": sc, "Soft Contingency": soft_contingency, "Carry": tcarry,
                       "Sales Cost": tsc, "Exit Value": rev, "Total Cost": total, "Profit": profit}
                clean_label = row_label.replace("**", "")
                row[f"${bc}/sf"] = f"${val[clean_label]:,.0f}"
            pf_data.append(row)
        st.dataframe(pf_data, use_container_width=True, hide_index=True)

        st.markdown("---")
        st.subheader("Exit Scenario Matrix")
        # Build scenario table
        test_psfs = sorted(set([
            int(breakeven_psf) if breakeven_psf > 0 else 300,
            median_psf - 25 if median_psf > 0 else 325,
            median_psf if median_psf > 0 else 350,
            median_psf + 25 if median_psf > 0 else 375,
            median_psf + 50 if median_psf > 0 else 400,
            exit_psf,
        ]))
        test_psfs = [p for p in test_psfs if p > 0]

        scenario_data = []
        for psf in test_psfs:
            revenue = psf * build_sf
            exit_c = revenue * (exit_cost_pct / 100)
            sc_cost = total_project_cost + total_interest + exit_c
            profit = revenue - sc_cost + net_rental_income
            margin = (profit / sc_cost) * 100 if sc_cost > 0 else 0
            label = ""
            if median_psf > 0 and psf == median_psf:
                label = "← MARKET MEDIAN"
            elif psf == exit_psf:
                label = "← YOUR EXIT"
            elif psf == int(breakeven_psf):
                label = "← BREAK-EVEN"
            scenario_data.append({
                "Exit $/sf": f"${psf}",
                "Revenue": f"${revenue:,.0f}",
                "Profit": f"${profit:,.0f}",
                "Margin": f"{margin:.1f}%",
                "Signal": "✅" if margin > 10 else ("⚠" if margin > 0 else "❌"),
                "Note": label,
            })
        st.dataframe(scenario_data, use_container_width=True, hide_index=True)

        # Sensitivity charts
        st.subheader("📈 Sensitivity — Purchase Price vs Profit")
        price_range = np.arange(max(200000, purchase_price - 200000),
                                purchase_price + 200000, 25000)
        profits_by_price = []
        for p in price_range:
            cost = p + total_dev_cost + total_interest + (exit_psf * build_sf * exit_cost_pct / 100)
            profits_by_price.append(exit_psf * build_sf - cost + net_rental_income)
        st.line_chart({"Purchase Price": price_range, "Profit": profits_by_price},
                      x="Purchase Price", y="Profit")

        st.subheader("📊 Sensitivity — Exit $/sf vs Profit")
        exit_range = np.arange(max(200, int(breakeven_psf) - 100), int(breakeven_psf) + 200, 10)
        profits_by_exit = []
        for e in exit_range:
            rev = e * build_sf
            exit_c = rev * (exit_cost_pct / 100)
            profits_by_exit.append(rev - (total_project_cost + total_interest + exit_c) + net_rental_income)
        st.line_chart({"Exit $/sf": exit_range, "Profit": profits_by_exit},
                      x="Exit $/sf", y="Profit")

        if median_psf > 0:
            st.info(f"📍 Market median is **${median_psf}/sf** — your break-even is **${breakeven_psf:.0f}/sf**. "
                    f"You need the market to be **${breakeven_psf - median_psf:+.0f}/sf above median** to break even.")

        # ── Carrying Cost Breakdown (During Construction) ──
        st.markdown("---")
        st.subheader("🏗️ Carrying Cost Breakdown (Construction Period)")
        st.caption(f"Build duration: {carry_months} months ({build_months} build + {delay_months} delays)")
        carry_data = [
            {"Component": "Loan Interest", "Amount": f"${carry_loan_interest:,.0f}",
             "Assumption": f"(Land+Dev) × {ltv}% LTC × {interest_rate}% × {draw_factor}% draw × {carry_months} mo ÷ 12"},
            {"Component": "Property Taxes", "Amount": f"${carry_taxes:,.0f}",
             "Assumption": f"(Land + 50% improvements) × {const_tax_rate}% × {carry_months} mo ÷ 12"},
            {"Component": "Insurance", "Amount": f"${carry_insurance:,.0f}",
             "Assumption": f"${const_insurance_annual:,}/yr × {carry_months} mo ÷ 12"},
            {"Component": "Utilities", "Amount": f"${carry_utilities:,.0f}",
             "Assumption": f"${const_utilities:,}/mo × {carry_months} mo"},
            {"Component": "Misc", "Amount": f"${carry_misc:,.0f}",
             "Assumption": f"${const_misc:,}/mo × {carry_months} mo"},
            {"Component": f"Buffer ({carry_buffer_pct}%)", "Amount": f"${carry_buffer:,.0f}",
             "Assumption": f"{carry_buffer_pct}% of carry subtotal"},
            {"Component": "**TOTAL CARRY**", "Amount": f"**${total_carry:,.0f}**", "Assumption": ""},
        ]
        st.dataframe(carry_data, use_container_width=True, hide_index=True)

        # ── Sales Cost Breakdown ──
        st.markdown("---")
        st.subheader("💸 Sales Cost Breakdown")
        st.caption(f"Based on exit value of ${user_revenue_est:,.0f} ({exit_psf}/sf × {build_sf:,} sf)")
        sales_data = [
            {"Component": f"Realtor ({broker_fee_pct}%)", "Amount": f"${user_revenue_est * broker_fee_pct / 100:,.0f}"},
            {"Component": f"Title + Closing ({title_closing_pct}%)", "Amount": f"${user_revenue_est * title_closing_pct / 100:,.0f}"},
            {"Component": f"Concessions ({seller_concessions_pct}%)", "Amount": f"${user_revenue_est * seller_concessions_pct / 100:,.0f}"},
            {"Component": f"Staging", "Amount": f"${staging_cost:,.0f}"},
            {"Component": f"Marketing", "Amount": f"${marketing_cost:,.0f}"},
            {"Component": f"Warranty", "Amount": f"${warranty_cost:,.0f}"},
            {"Component": f"Holding During Sale ({sale_hold_months} mo)", "Amount": f"${holding_during_sale:,.0f}"},
            {"Component": "**TOTAL SALES COST**", "Amount": f"**${total_sales_cost:,.0f}**"},
        ]
        st.dataframe(sales_data, use_container_width=True, hide_index=True)

        # ── Soft Cost Breakdown ──
        if split_soft:
            st.markdown("---")
            st.subheader("📐 Soft Cost Breakdown")
            soft_data = [
                {"Category": f"Architecture & Design ({arch_pct}%)", "Amount": f"${hard_cost * arch_pct / 100:,.0f}"},
                {"Category": f"Engineering ({eng_pct}%)", "Amount": f"${hard_cost * eng_pct / 100:,.0f}"},
                {"Category": f"Permits & Impact Fees ({permit_fee_pct}%)", "Amount": f"${hard_cost * permit_fee_pct / 100:,.0f}"},
                {"Category": f"Surveys & Geotech ({survey_pct}%)", "Amount": f"${hard_cost * survey_pct / 100:,.0f}"},
                {"Category": f"Builder's Risk Insurance ({insurance_dev_pct}%)", "Amount": f"${hard_cost * insurance_dev_pct / 100:,.0f}"},
                {"Category": f"Other Soft ({other_soft_pct}%)", "Amount": f"${hard_cost * other_soft_pct / 100:,.0f}"},
                {"Category": f"**TOTAL SOFT ({soft_cost_pct:.1f}%)**", "Amount": f"**${soft_costs:,.0f}**"},
            ]
            st.dataframe(soft_data, use_container_width=True, hide_index=True)

        # ── Cost vs Price Profit Matrix ──
        st.markdown("---")
        st.subheader("📊 Cost vs Price Profit Matrix")
        st.caption("Profit at different build cost and exit price combinations")
        matrix_build_costs = sorted(set([max(100, build_cost_psf - 50), max(125, build_cost_psf - 25), build_cost_psf, build_cost_psf + 25, build_cost_psf + 50]))
        matrix_exit_prices = sorted(set([max(300, exit_psf - 75), max(325, exit_psf - 50), max(350, exit_psf - 25), exit_psf, exit_psf + 25, exit_psf + 50]))

        matrix_data = []
        for bc in matrix_build_costs:
            row = {"Build Cost": f"${bc}/sf"}
            for ep in matrix_exit_prices:
                hc = bc * build_sf
                hc_cont = hc * (hard_contingency_pct / 100)
                sc = hc * (soft_cost_pct / 100)
                dev = hc + hc_cont + sc + soft_contingency
                proj = purchase_price + dev
                loan = dev * (ltv / 100)
                ci = loan * (interest_rate / 100) * (draw_factor / 100) * carry_months / 12
                lf = loan * (loan_fee_pct / 100)
                # Carry costs
                ct = (purchase_price + 0.5 * dev) * (const_tax_rate / 100) * carry_months / 12
                cins = const_insurance_annual * carry_months / 12
                cutil = const_utilities * carry_months
                cmisc = const_misc * carry_months
                csub = ci + ct + cins + cutil + cmisc
                cbuf = csub * (carry_buffer_pct / 100)
                tcarry = csub + cbuf
                rev = ep * build_sf
                ec = rev * (exit_cost_pct / 100)
                stg = staging_base + staging_per_unit * units
                mkt = marketing_base + marketing_per_unit * units
                war = warranty_per_unit * units
                hds = tcarry / max(carry_months, 1) * sale_hold_months
                tsc = ec + stg + mkt + war + hds
                total = purchase_price + dev + tcarry + tsc + lf
                profit = rev - total
                row[f"${ep}/sf"] = f"${profit:,.0f}"
            matrix_data.append(row)
        st.dataframe(matrix_data, use_container_width=True, hide_index=True)

        # ── Timeline Sensitivity ──
        st.markdown("---")
        st.subheader("⏱️ Timeline Sensitivity — Profit by Duration")
        st.caption(f"Exit at ${exit_psf}/sf, build cost ${build_cost_psf}/sf")
        timeline_durations = list(range(max(6, build_months - 3), build_months + 5))
        cost_scenarios = sorted(set([max(100, build_cost_psf - 50), max(125, build_cost_psf - 25), build_cost_psf, build_cost_psf + 25, build_cost_psf + 50]))

        timeline_data = []
        for dur in timeline_durations:
            row = {"Duration": f"{dur} months"}
            for bc in cost_scenarios:
                hc = bc * build_sf
                hc_cont = hc * (hard_contingency_pct / 100)
                sc = hc * (soft_cost_pct / 100)
                dev = hc + hc_cont + sc + soft_contingency
                proj = purchase_price + dev
                loan = dev * (ltv / 100)
                ci = loan * (interest_rate / 100) * (draw_factor / 100) * dur / 12
                lf = loan * (loan_fee_pct / 100)
                ct = (purchase_price + 0.5 * dev) * (const_tax_rate / 100) * dur / 12
                cins = const_insurance_annual * dur / 12
                cutil = const_utilities * dur
                cmisc = const_misc * dur
                csub = ci + ct + cins + cutil + cmisc
                cbuf = csub * (carry_buffer_pct / 100)
                tcarry = csub + cbuf
                rev = exit_psf * build_sf
                ec = rev * (exit_cost_pct / 100)
                stg = staging_base + staging_per_unit * units
                mkt = marketing_base + marketing_per_unit * units
                war = warranty_per_unit * units
                hds = tcarry / max(dur, 1) * sale_hold_months
                tsc = ec + stg + mkt + war + hds
                total = purchase_price + dev + tcarry + tsc + lf
                profit = rev - total
                row[f"${bc}/sf"] = f"${profit:,.0f}"
            timeline_data.append(row)
        st.dataframe(timeline_data, use_container_width=True, hide_index=True)

        # ── Rent Fallback / Hold Check ──
        st.markdown("---")
        st.subheader("🏠 Rent Fallback / Hold Check")
        rent_gross_monthly = rent_per_unit * units
        rent_gross_annual = rent_gross_monthly * 12
        # Use $250/sf scenario total cost as baseline
        baseline_total = purchase_price + total_dev_cost + total_carry + total_sales_cost + loan_fees
        rent_yield = (rent_gross_annual / baseline_total * 100) if baseline_total > 0 else 0
        rent_vs_exit = (rent_gross_annual / user_revenue_est * 100) if user_revenue_est > 0 else 0

        if rent_yield >= 5.5:
            rent_read = "✅ Reasonable fallback"
        elif rent_yield >= 4.5:
            rent_read = "⚠️ Thin fallback"
        else:
            rent_read = "❌ Weak rental coverage"

        rf1, rf2, rf3, rf4 = st.columns(4)
        rf1.metric("Units", units)
        rf2.metric("Rent / Unit / Mo", f"${rent_per_unit:,}")
        rf3.metric("Gross Annual Rent", f"${rent_gross_annual:,}")
        rf4.metric("Rent Yield", f"{rent_yield:.1f}%", delta=rent_read)

        st.markdown(f"""
        | Metric | Value | Note |
        |--------|-------|------|
        | Gross Monthly Rent | ${rent_gross_monthly:,} | {units} units × ${rent_per_unit:,}/mo |
        | Gross Annual Rent | ${rent_gross_annual:,} | Monthly × 12 |
        | Annual Rent / Total Cost | {rent_yield:.1f}% | Gross rent ÷ all-in cost |
        | Annual Rent / Exit Value | {rent_vs_exit:.1f}% | Gross rent ÷ exit value |
        | **Assessment** | **{rent_read}** | ≥5.5% = Reasonable, ≥4.5% = Thin |
        """)

    # ── Comps Tab ──
    with tab_comps:
        if result.redfin_comps:
            # Similar-size comps (per unit)
            sim_stats = getattr(result, 'similar_stats', {})
            if sim_stats:
                st.subheader(f"🎯 Similar Size Comps ({sim_stats['size_range']}, per unit = {int(per_unit_sf):,} sf)")
                sc1, sc2, sc3, sc4 = st.columns(4)
                sc1.metric("Median $/sf", f"${sim_stats.get('median_psf', 0)}")
                sc2.metric("Average $/sf", f"${sim_stats.get('avg_psf', 0)}")
                sc3.metric("Count", f"{sim_stats.get('count', 0)}")
                sc4.metric("Range", f"${sim_stats.get('min_psf', 0)}–${sim_stats.get('max_psf', 0)}")

                sim_data = []
                for c in sorted(similar_comps, key=lambda x: x.get("psf", 0), reverse=True):
                    if c.get("psf", 0) > 0:
                        sim_data.append({
                            "Address": c["address"],
                            "Price": f"${c['price']:,}",
                            "Size": f"{c['sqft']:,} sf",
                            "$/sf": c["psf"],
                            "Built": c.get("year_built", ""),
                            "Sold": c.get("sold_date", ""),
                            "Redfin": c.get("redfin_url", ""),
                            "Zillow": c.get("zillow_url", ""),
                        })
                st.dataframe(sim_data, use_container_width=True, hide_index=True,
                             column_config={
                                 "Redfin": st.column_config.LinkColumn("Redfin", display_text="View"),
                                 "Zillow": st.column_config.LinkColumn("Zillow", display_text="View"),
                             })
                st.markdown("---")

            # All comps
            stats = result.market_stats
            if not result.sources_status.get('comps_radius', True):
                st.warning("⚠️ Could not pull comps within 1 mile — showing ZIP-wide data as fallback.")
                st.subheader(f"📍 Sold Comps — ZIP {zip_code} (past 2 years, {stats.get('count', 0)} comps)")
            else:
                st.subheader(f"📍 Sold Comps — Within 1 Mile (past 2 years, {stats.get('count', 0)} comps)")
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Median $/sf", f"${stats.get('median_psf', 0)}")
            c2.metric("Average $/sf", f"${stats.get('avg_psf', 0)}")
            c3.metric("Min $/sf", f"${stats.get('min_psf', 0)}")
            c4.metric("Max $/sf", f"${stats.get('max_psf', 0)}")

            comp_data = []
            for c in sorted(result.redfin_comps, key=lambda x: x.get("distance_mi", 0)):
                if c.get("psf", 0) > 0:
                    comp_data.append({
                        "Address": c["address"],
                        "Price": f"${c['price']:,}",
                        "Size": f"{c['sqft']:,} sf",
                        "$/sf": c["psf"],
                        "Distance": f"{c.get('distance_mi', 0):.1f} mi",
                        "Built": c.get("year_built", ""),
                        "Sold": c.get("sold_date", ""),
                        "Beds": c.get("beds", ""),
                        "Baths": c.get("baths", ""),
                        "Redfin": c.get("redfin_url", ""),
                        "Zillow": c.get("zillow_url", ""),
                    })
            st.dataframe(comp_data, use_container_width=True, hide_index=True,
                         column_config={
                             "Redfin": st.column_config.LinkColumn("Redfin", display_text="View"),
                             "Zillow": st.column_config.LinkColumn("Zillow", display_text="View"),
                         })
        else:
            if result.market_stats and result.market_stats.get('source', '').startswith('Census'):
                st.warning("⚠️ Redfin blocked — using **Census Bureau estimate** as fallback.")
                stats = result.market_stats
                st.subheader(f"📊 Estimated Market Value — ZIP {zip_code}")
                c1, c2 = st.columns(2)
                c1.metric("Estimated $/sf", f"${stats.get('median_psf', 0)}")
                c2.metric("Source", "Census ACS 2022")
                st.caption("Based on Census median home value ÷ typical new construction size. For accurate comps, check Redfin/Zillow directly.")
            else:
                st.warning("No Redfin comps available. Redfin may be blocking requests from this server.")

    # ── Active Listings Tab ──
    with tab_active:
        st.subheader(f"🏠 Active Listings — Within 1 Mile")
        if result.active_comps:
            a_stats = result.active_stats
            if a_stats:
                ac1, ac2, ac3, ac4 = st.columns(4)
                ac1.metric("Median $/sf", f"${a_stats.get('median_psf', 0)}")
                ac2.metric("Average $/sf", f"${a_stats.get('avg_psf', 0)}")
                ac3.metric("Count", f"{a_stats.get('count', 0)}")
                ac4.metric("Range", f"${a_stats.get('min_psf', 0)}–${a_stats.get('max_psf', 0)}")

            a_data = []
            for c in result.active_comps:
                if c.get("psf", 0) > 0:
                    a_data.append({
                        "Address": c["address"],
                        "List Price": f"${c['price']:,}",
                        "Size": f"{c['sqft']:,} sf",
                        "$/sf": c["psf"],
                        "Distance": f"{c.get('distance_mi', 0):.1f} mi",
                        "Built": c.get("year_built", ""),
                        "Beds": c.get("beds", ""),
                        "Baths": c.get("baths", ""),
                        "Days on Market": c.get("days_on_market", ""),
                        "Redfin": c.get("redfin_url", ""),
                        "Zillow": c.get("zillow_url", ""),
                    })
            st.dataframe(a_data, use_container_width=True, hide_index=True,
                         column_config={
                             "Redfin": st.column_config.LinkColumn("Redfin", display_text="View"),
                             "Zillow": st.column_config.LinkColumn("Zillow", display_text="View"),
                         })
            st.caption("💡 These are your competition — currently listed homes near the subject property.")
        else:
            st.info("No active listings found within 1 mile.")

    # ── Rental Comps Tab ──
    with tab_rental:
        st.subheader("💰 Rental Market Analysis")

        # ── Research Links (at top for easy access) ──
        st.markdown("### 🔍 Search Active Rentals Near You")
        zillow_rental = f"https://www.zillow.com/homes/for_rent/{zip_code}_rb/"
        apartments_url = f"https://www.apartments.com/{zip_code}/"
        rentometer_url = f"https://www.rentometer.com/analysis/new?address={zip_code}"
        redfin_rental = f"https://www.redfin.com/zipcode/{zip_code}/apartments-for-rent"

        lk1, lk2, lk3, lk4 = st.columns(4)
        lk1.markdown(f"[🏠 Zillow Rentals]({zillow_rental})")
        lk2.markdown(f"[🏢 Apartments.com]({apartments_url})")
        lk3.markdown(f"[📊 Rentometer]({rentometer_url})")
        lk4.markdown(f"[🔴 Redfin Rentals]({redfin_rental})")
        st.caption("Click above to search active rental listings in your ZIP code.")

        st.markdown("---")

        # ── Census Bureau Rental Data ──
        census_rents = fetch_census_rents(zip_code)
        if census_rents and census_rents.get('median_rent'):
            st.markdown(f"### 📊 Actual Median Rents — ZIP {zip_code}")
            st.caption(f"Source: {census_rents.get('source', 'Census ACS')}")

            cr1, cr2, cr3 = st.columns(3)
            cr1.metric("Overall Median Rent", f"${census_rents['median_rent']:,}/mo")
            beds_in_unit = max(1, int(build_sf / max(units, 1) / 500))  # rough estimate
            cr2.metric("Your Unit Size (~" + str(beds_in_unit) + "BR)", 
                       f"${census_rents.get(f'rent_{beds_in_unit}br', census_rents['median_rent']) or census_rents['median_rent']:,}/mo")
            cr3.metric("Your Assumption", f"${rent_per_unit:,}/mo",
                       delta=f"{((rent_per_unit / census_rents['median_rent']) - 1) * 100:+.0f}% vs median" if census_rents['median_rent'] else None)

            # Bedroom breakdown table
            bed_data = []
            for label, key in [("Studio", "rent_0br"), ("1 Bedroom", "rent_1br"), ("2 Bedroom", "rent_2br"), 
                               ("3 Bedroom", "rent_3br"), ("4+ Bedroom", "rent_4br_plus")]:
                val = census_rents.get(key)
                if val:
                    bed_data.append({"Type": label, "Median Rent": f"${val:,}/mo", 
                                    "Annual": f"${val * 12:,}",
                                    "vs Your Assumption": f"{((rent_per_unit / val) - 1) * 100:+.1f}%"})
            if bed_data:
                st.dataframe(bed_data, use_container_width=True, hide_index=True)

            st.markdown("---")

        # ── a) Your Rent Assumptions ──
        gross_monthly_rent = rent_per_unit * units
        gross_annual_rent = gross_monthly_rent * 12
        rent_per_sf = gross_monthly_rent / build_sf if build_sf > 0 else 0
        gross_yield = (gross_annual_rent / total_project_cost * 100) if total_project_cost > 0 else 0

        st.markdown("### 📋 Your Rent Assumptions")
        rc1, rc2, rc3, rc4 = st.columns(4)
        rc1.metric("Gross Monthly Rent", f"${gross_monthly_rent:,.0f}")
        rc2.metric("Annual Rent", f"${gross_annual_rent:,.0f}")
        rc3.metric("Rent/SF/mo", f"${rent_per_sf:.2f}")
        rc4.metric("Gross Yield", f"{gross_yield:.1f}%")

        # ── b) Market Rent Estimates ──
        st.markdown("### 📊 Market Rent Estimates")

        # 1% rule estimate
        one_pct_monthly = total_project_cost * 0.01
        one_pct_annual = one_pct_monthly * 12
        one_pct_per_sf = one_pct_monthly / build_sf if build_sf > 0 else 0

        # Cap rate estimates from sold comps
        estimated_value = median_psf * build_sf if median_psf > 0 else total_project_cost
        cap5_annual = estimated_value * 0.05
        cap5_monthly = cap5_annual / 12
        cap5_per_sf = cap5_monthly / build_sf if build_sf > 0 else 0

        cap7_annual = estimated_value * 0.07
        cap7_monthly = cap7_annual / 12
        cap7_per_sf = cap7_monthly / build_sf if build_sf > 0 else 0

        comparison_data = [
            {"Method": "Your Assumption", "Monthly Rent": f"${gross_monthly_rent:,.0f}", "Annual": f"${gross_annual_rent:,.0f}", "$/SF/mo": f"${rent_per_sf:.2f}"},
        ]
        if census_rents and census_rents.get('median_rent'):
            cm = census_rents['median_rent'] * units
            comparison_data.append({"Method": f"Census Median (×{units} units)", "Monthly Rent": f"${cm:,.0f}", "Annual": f"${cm * 12:,.0f}", "$/SF/mo": f"${cm / build_sf:.2f}" if build_sf > 0 else "—"})
        comparison_data.extend([
            {"Method": "1% Rule (of cost)", "Monthly Rent": f"${one_pct_monthly:,.0f}", "Annual": f"${one_pct_annual:,.0f}", "$/SF/mo": f"${one_pct_per_sf:.2f}"},
            {"Method": "5% Cap Rate", "Monthly Rent": f"${cap5_monthly:,.0f}", "Annual": f"${cap5_annual:,.0f}", "$/SF/mo": f"${cap5_per_sf:.2f}"},
            {"Method": "7% Cap Rate", "Monthly Rent": f"${cap7_monthly:,.0f}", "Annual": f"${cap7_annual:,.0f}", "$/SF/mo": f"${cap7_per_sf:.2f}"},
        ])
        st.dataframe(comparison_data, use_container_width=True, hide_index=True)

        if median_psf > 0:
            st.caption(f"💡 Cap rate estimates based on comps median ${median_psf}/sf × {build_sf:,} sf = ${estimated_value:,.0f} estimated value")
        else:
            st.caption("💡 Cap rate estimates based on total project cost (no comp data available)")

        # ── c) Rent Sensitivity Analysis ──
        st.markdown("### 📈 Rent Sensitivity Analysis")

        sensitivity_data = []
        for pct in [-20, -10, 0, 10, 20]:
            adj_rent_monthly = gross_monthly_rent * (1 + pct / 100)
            adj_rent_annual = adj_rent_monthly * 12

            # Recalc NOI and CF for this rent level
            eff_rent_mo = adj_rent_monthly * (1 - vacancy_pct / 100)
            mgmt_mo = eff_rent_mo * (mgmt_fee_pct / 100)
            tax_mo = build_sf * taxable_value_psf * (prop_tax_rate / 100) / 12
            sens_monthly_noi = eff_rent_mo - mgmt_mo - tax_mo - insurance_monthly - (repairs_per_unit * units) - common_utilities - leasing_reserve
            sens_monthly_cf = sens_monthly_noi - monthly_debt_service
            sens_annual_cf = sens_monthly_cf * 12
            sens_coc = (sens_annual_cf / total_equity_invested * 100) if total_equity_invested > 0 else 0

            label = f"{pct:+d}%" if pct != 0 else "0% (Base)"
            sensitivity_data.append({
                "Scenario": label,
                "Monthly Rent": f"${adj_rent_monthly:,.0f}",
                "Monthly NOI": f"${sens_monthly_noi:,.0f}",
                "Monthly CF After Debt": f"${sens_monthly_cf:,.0f}",
                "Annual CF": f"${sens_annual_cf:,.0f}",
                "Cash-on-Cash": f"{sens_coc:.1f}%",
            })

        st.dataframe(sensitivity_data, use_container_width=True, hide_index=True)
        st.caption("💡 The **0% (Base)** row matches your assumed rent. Negative CF means you'd need to fund the shortfall from cash reserves.")

    # ── Permits Tab ──
    with tab_permits:
        st.subheader(f"New Construction Permits — {street_name} St")
        if result.street_permits:
            active = [p for p in result.street_permits if p.status == 'Active']
            final = [p for p in result.street_permits if p.status == 'Final']

            if active:
                st.markdown(f"### 🟡 Under Construction ({len(active)})")
                st.dataframe([{"Address": p.address, "Size": f"{p.sqft:,.0f} sf",
                              "Builder": p.builder, "Contractor": p.contractor_name,
                              "Applicant": p.applicant_name or p.applicant_org,
                              "Date": p.issue_date,
                              "Type": p.permit_class,
                              "Source": f"https://data.austintexas.gov/resource/3syk-w9eu.json?permit_num={p.permit_number}" if p.permit_number else ""} for p in active],
                             use_container_width=True, hide_index=True,
                             column_config={"Source": st.column_config.LinkColumn("Source", display_text="Austin Open Data")})
            if final:
                st.markdown(f"### ✅ Completed ({len(final)})")
                st.dataframe([{"Address": p.address, "Size": f"{p.sqft:,.0f} sf",
                              "Builder": p.builder, "Contractor": p.contractor_name,
                              "Applicant": p.applicant_name or p.applicant_org,
                              "Date": p.issue_date,
                              "Type": p.permit_class,
                              "Source": f"https://data.austintexas.gov/resource/3syk-w9eu.json?permit_num={p.permit_number}" if p.permit_number else ""} for p in final],
                             use_container_width=True, hide_index=True,
                             column_config={"Source": st.column_config.LinkColumn("Source", display_text="Austin Open Data")})

            # Zip summary
            st.markdown(f"### 📊 Zip {zip_code} — Permit Trend")
            by_year = {}
            for p in result.zip_permits:
                yr = p.issue_date[:4] if p.issue_date else 'Unknown'
                by_year.setdefault(yr, []).append(p)
            year_data = []
            for yr in sorted(by_year.keys(), reverse=True):
                permits = by_year[yr]
                total_sf = sum(p.sqft for p in permits)
                year_data.append({"Year": yr, "Permits": len(permits),
                                  "Total SF": f"{total_sf:,.0f}",
                                  "Avg SF": f"{total_sf / len(permits):,.0f}" if permits else "0"})
            st.dataframe(year_data, use_container_width=True, hide_index=True)
        else:
            st.warning("No new construction permits found. This could mean:\n"
                      "- No recent builds on this street\n"
                      "- The permits API may be temporarily unavailable — try clicking **Analyze** again")

        # All permit types for the street
        st.markdown("---")
        st.subheader(f"📋 All Permits — {street_name} St (past 2 years)")
        if result.street_all_permits:
            # Group by type
            type_map = {'BP': '🏗️ Building', 'EP': '⚡ Electrical', 'PP': '🔧 Plumbing',
                        'MP': '🌬️ Mechanical', 'DP': '🚧 Demolition'}
            by_type = {}
            for p in result.street_all_permits:
                key = p.permit_type or 'Other'
                by_type.setdefault(key, []).append(p)

            # Summary counts
            type_cols = st.columns(min(len(by_type), 5))
            for i, (ptype, permits) in enumerate(sorted(by_type.items())):
                label = type_map.get(ptype, ptype)
                type_cols[i % len(type_cols)].metric(label, len(permits))

            # Table with all permits
            all_data = []
            for p in result.street_all_permits:
                all_data.append({
                    "Address": p.address,
                    "Category": type_map.get(p.permit_type, p.permit_type),
                    "Work": p.work_class,
                    "Description": (p.description or "")[:120],
                    "Submitted By": p.applicant_name or p.applicant_org or "",
                    "Contractor": p.builder if p.builder != "Unknown" else (p.contractor_name or ""),
                    "Status": p.status,
                    "Date": p.issue_date,
                })
            st.dataframe(all_data, use_container_width=True, hide_index=True)
        else:
            st.info("No permits found for this street in the past 2 years.")

    # ── Plot Info Tab ──
    with tab_plot:
        st.subheader(f"📋 Plot Information — {address}")

        # Listing status banner
        listing = result.listing_status
        listing_status = listing.get('status', 'Unknown')
        listing_url = listing.get('url', '')
        listing_price = listing.get('price', '')

        if listing_status == 'Active':
            st.success(f"🟢 **ACTIVE LISTING** — This property is for sale! "
                       f"{'Listed at $' + f'{int(listing_price):,}' if listing_price else ''} "
                       f"{'[View on Redfin →](' + listing_url + ')' if listing_url else ''}")
        elif listing_status == 'Pending':
            st.warning(f"🟡 **PENDING / UNDER CONTRACT** — Someone already has an offer accepted on this property. "
                       f"{'Listed at $' + f'{int(listing_price):,}' if listing_price else ''} "
                       f"You'd need to wait for it to fall through, or find a similar property. "
                       f"{'[View on Redfin →](' + listing_url + ')' if listing_url else ''}")
        elif listing_status == 'Sold':
            st.error(f"🔴 **SOLD** — This property has already been sold. "
                     f"{'Sale price: $' + f'{int(listing_price):,}' if listing_price else ''} "
                     f"{'[View on Redfin →](' + listing_url + ')' if listing_url else ''}")
        elif listing_status == 'Not Found':
            st.info("⚪ **OFF-MARKET** — This property is not currently listed on Redfin. "
                    "It may be a private sale, pocket listing, or not yet on the market.")
        else:
            st.info("⚪ **Listing status unknown** — Could not verify on Redfin.")

        # Property links — use Redfin URL from listing check if available
        addr_slug = address.replace(' ', '-').replace(',', '')
        addr_query = address.replace(' ', '+')
        zillow_search = f"https://www.zillow.com/homes/{addr_slug}-Austin-TX-{zip_code}_rb/"
        redfin_link = listing_url if listing_url else f"https://www.google.com/search?q=site:redfin.com+{addr_query}+Austin+TX+{zip_code}"
        tcad_search = f"https://stage.travis.prodigycad.com/property-search"
        google_maps = f"https://www.google.com/maps/search/{addr_query}+Austin+TX+{zip_code}"

        lc1, lc2, lc3, lc4 = st.columns(4)
        lc1.markdown(f"[🔗 Zillow]({zillow_search})")
        lc2.markdown(f"[🔗 Redfin]({redfin_link})")
        lc3.markdown(f"[🔗 TCAD]({tcad_search})")
        lc4.markdown(f"[🗺️ Google Maps]({google_maps})")

        st.markdown("---")

        # Geocode and fetch plot data
        with st.spinner("Fetching zoning, parcel & flood data..."):
            lat, lon = geocode_address(address, zip_code)

        if lat and lon:
            plot_data = fetch_plot_info(lat, lon, address)
            st.caption(f"📍 Coordinates: {lat:.6f}, {lon:.6f}")

            pc1, pc2 = st.columns(2)

            # Zoning
            with pc1:
                st.markdown("### 🏗️ Zoning")
                zoning = plot_data.get('zoning', {})
                if zoning:
                    ztype = zoning.get('zoning_type', 'Unknown')
                    base = zoning.get('base_zone', '')
                    name = zoning.get('zone_name', '')
                    zoning_area = zoning.get('lot_area_sf', 0)

                    st.metric("Zoning Type", ztype)

                    # Plain-English explanation
                    info = ZONING_INFO.get(base, {})
                    if info:
                        st.info(f"💡 **What this means:** {info.get('plain', '')}")

                        st.markdown(f"""
                        | | Detail |
                        |---|---|
                        | **Official Name** | {name} |
                        | **What you can build** | {info.get('can_build', 'Check with city')} |
                        | **Max height** | {info.get('height', 'Check with city')} |
                        | **Max units allowed** | {info.get('max_units', 'Unknown')} |
                        | **Min lot size required** | {info.get('min_lot_sf', 0):,} sf |
                        | **This lot size** | {zoning_area:,} sf ({zoning_area/43560:.2f} acres) |
                        """)

                        if zoning_area > 0 and isinstance(info.get('max_units'), str) and '/acre' in str(info['max_units']):
                            density = int(info['max_units'].replace('/acre', ''))
                            acres = zoning_area / 43560
                            est_units = int(acres * density)
                            st.success(f"📐 **Estimated max units on this lot:** ~{est_units} ({acres:.2f} acres × {density}/acre)")
                    else:
                        st.write(f"**Full Name:** {name}")
                        st.write(f"**Base Zone:** {base}")
                        if zoning_area > 0:
                            st.write(f"**Lot Area:** {zoning_area:,} sf ({zoning_area/43560:.2f} acres)")

                    # Overlay explanations
                    for suffix, (title, explanation) in OVERLAY_EXPLANATIONS.items():
                        if suffix in ztype:
                            st.markdown("---")
                            st.markdown(f"#### {title}")
                            st.write(explanation)
                else:
                    st.warning("Zoning data not available for this location")

            # Parcel & Flood
            with pc2:
                st.markdown("### 📐 Parcel")
                parcel = plot_data.get('parcel', {})
                if parcel:
                    parcel_sf = parcel.get('parcel_area_sf', 0)
                    st.write(f"**TCAD Property ID:** {parcel.get('prop_id', 'N/A')}")
                    st.write(f"**Parcel ID:** {parcel.get('pid', 'N/A')}")
                    st.write(f"**Lot:** {parcel.get('lot', 'N/A')}, **Block:** {parcel.get('block', 'N/A')}")
                    if parcel_sf > 0:
                        st.write(f"**Parcel Area:** {parcel_sf:,} sf ({parcel_sf/43560:.3f} acres)")
                    tcad_link = f"https://stage.travis.prodigycad.com/property-detail/{parcel.get('prop_id', '')}/2026"
                    st.markdown(f"[View on TCAD →]({tcad_link})")
                else:
                    st.warning("Parcel data not available")

                st.markdown("### 🌊 FEMA Flood Zone")
                flood = plot_data.get('flood', {})
                if flood.get('in_floodplain'):
                    st.error(f"❌ **IN FLOOD ZONE:** {flood.get('zone', 'Unknown')}")
                    st.write("**What this means:** Your property is in a flood-risk area. "
                             "You'll be **required to buy flood insurance** (can be $1,000–$5,000+/year). "
                             "Construction costs may be higher (elevated foundation). "
                             "Resale can be harder — many buyers avoid flood zones.")
                else:
                    st.success(f"✅ **Not in floodplain** — Zone: {flood.get('zone', 'X')}")
                    st.write("**What this means:** Low flood risk. No mandatory flood insurance required. "
                             "This is the best-case scenario for lenders and insurance costs.")

                st.markdown("### 🛣️ Legal Access to Road")
                st.info("**What this means:** Every buildable lot needs legal access to a public road. "
                        "If the lot is landlocked (no road frontage), you can't get a building permit. "
                        "Check the **plat map** to confirm the lot touches a public street or has a recorded easement.")
                st.markdown(f"[🗺️ View Austin GIS Map](https://www.austintexas.gov/gis/) · "
                           f"[📄 Travis County Plat Records](https://www.traviscountyclerk.org/)")

                st.markdown("### 📜 Title & Liens")
                st.info("**What this means:** Before buying, a title search checks if anyone else "
                        "has a legal claim on the property — unpaid taxes, mortgages, contractor liens, "
                        "or easements. Your title company does this at closing, but you can check early.")
                st.markdown("[🔍 Travis County Deed Records](https://deed.traviscountyclerk.org/)")

        else:
            st.error("Could not geocode address. Please verify the address and ZIP code.")

    # ── AI Analysis Tab ──
    with tab_ai:
        st.subheader("🤖 AI Deal Analysis")
        try:
            if not _get_gemini_key():
                st.info("**How to enable free AI analysis:**\n\n"
                        "1. Go to [Google AI Studio](https://aistudio.google.com/apikey) and get a free API key\n"
                        "2. In Streamlit Cloud → Settings → Secrets, add:\n"
                        "```\nGEMINI_API_KEY = \"your-api-key-here\"\n```\n"
                        "3. Refresh the app — the 🤖 AI Analysis tab will work!\n\n"
                        "**Cost: FREE** (Google Gemini free tier: 15 requests/minute)")
            else:
                if st.button("🧠 Generate AI Analysis", type="primary"):
                    with st.spinner("AI is analyzing your deal..."):
                        # Build top comps summary
                        top_comps_text = ""
                        if result.redfin_comps:
                            for c in sorted(result.redfin_comps, key=lambda x: x.get('distance_mi', 0))[:5]:
                                top_comps_text += f"  - {c['address']}: ${c['price']:,} ({c['sqft']:,}sf, ${c['psf']}/sf, sold {c.get('sold_date', 'N/A')}, {c.get('distance_mi', 0):.1f}mi away)\n"

                        # Fetch Census rents for AI
                        _census_ai = fetch_census_rents(zip_code)
                        _census_median = f"${_census_ai.get('median_rent', 0):,}" if _census_ai and _census_ai.get('median_rent') else 'N/A'
                        _census_by_br = ', '.join([
                            f"{br}: ${_census_ai.get(k, 0):,}/mo"
                            for br, k in [('Studio', 'rent_0br'), ('1BR', 'rent_1br'), ('2BR', 'rent_2br'), ('3BR', 'rent_3br'), ('4BR+', 'rent_4br_plus')]
                            if _census_ai and _census_ai.get(k)
                        ]) if _census_ai else 'N/A'

                        deal_data = {
                            'address': address,
                            'zip_code': zip_code,
                            'purchase_price': purchase_price,
                            'total_sf': build_sf,
                            'num_units': units,
                            'per_unit_sf': per_unit_sf,
                            'exit_psf': exit_psf,
                            'total_cost': total_cost,
                            'revenue': adjusted_revenue,
                            'profit': market_profit,
                            'margin_pct': (market_profit / total_cost * 100) if total_cost > 0 else 0,
                            'breakeven_psf': breakeven_psf,
                            'median_psf': median_psf if median_psf > 0 else result.market_stats.get('median_psf', 0),
                            'comp_count': result.market_stats.get('count', 0) or len(result.redfin_comps),
                            'risk_score': risk_score,
                            'verdict': verdict,
                            'listing_status': result.listing_status.get('status', 'Unknown') if hasattr(result, 'listing_status') else 'Unknown',
                            'zoning': 'N/A',
                            'monthly_rent': rent_per_unit * units,
                            'rental_noi': monthly_noi * 12,
                            # Top 5 sold comps summary
                            'top_comps': top_comps_text,
                            # Permit activity
                            'active_permits': len([p for p in result.street_permits if p.status == 'Active']),
                            'completed_permits': len([p for p in result.street_permits if p.status == 'Final']),
                            'permit_builders': ', '.join(set(p.builder for p in result.street_permits[:5] if p.builder != 'Unknown')),
                            # Hold model financials
                            'hold_months': hold_months,
                            'monthly_noi': monthly_noi,
                            'monthly_cf_after_debt': monthly_cf_after_debt,
                            'monthly_debt_service': monthly_debt_service,
                            'equity_multiple': equity_multiple,
                            'total_equity_invested': total_equity_invested,
                            'annualized_return': annualized_return * 100,
                            'cash_yield': cash_yield,
                            # Construction
                            'build_months': build_months,
                            'construction_interest': construction_interest,
                            'loan_amount': loan_amount,
                            'build_cost_psf': build_cost_psf,
                            # Census rental data
                            'census_median_rent': _census_median,
                            'census_rent_by_br': _census_by_br,
                        }
                        ai_text = generate_ai_summary(deal_data)
                        if ai_text:
                            # Escape $ signs so Streamlit doesn't render as LaTeX math
                            ai_text = ai_text.replace('$', '\\$')
                            st.markdown(ai_text)
                        else:
                            st.error("Could not generate AI analysis. Check your API key.")
        except Exception as e:
            st.error(f"AI Analysis error: {e}")

    # ── Download Tab ──
    with tab_download:
        st.subheader("Download Report")
        if result is not None:
            report_bytes = generate_report_bytes(
                result, address, zip_code, street_name,
                purchase_price, build_sf, exit_psf, build_cost_psf
            )
            st.download_button(
                label="📥 Download Word Report (.docx)",
                data=report_bytes,
                file_name=f"{street_name.title()}_St_Analysis.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary",
            )

            # Excel download
            try:
                _nsbd = net_sale_before_debt if hold_months > 0 else (user_revenue - exit_costs_user)
                _prcf = positive_rental_cf if hold_months > 0 else 0
                _nsad = net_sale_after_debt if hold_months > 0 else (user_revenue - exit_costs_user - loan_balance_after_hold)
                excel_bytes = generate_excel_bytes(
                units=units, per_unit_sf=per_unit_sf, build_sf=build_sf,
                purchase_price=purchase_price, build_cost_psf=build_cost_psf, exit_psf=exit_psf,
                build_months=build_months, hold_months=hold_months, delay_months=delay_months,
                hard_contingency_pct=hard_contingency_pct, soft_cost_pct=soft_cost_pct,
                soft_contingency=soft_contingency,
                ltv=ltv, interest_rate=interest_rate, draw_factor=draw_factor,
                loan_fee_pct=loan_fee_pct,
                perm_mortgage_rate=perm_mortgage_rate, amortization_years=amortization_years,
                rent_per_unit=rent_per_unit, vacancy_pct=vacancy_pct, mgmt_fee_pct=mgmt_fee_pct,
                prop_tax_rate=prop_tax_rate, taxable_value_psf=taxable_value_psf,
                insurance_monthly=insurance_monthly,
                repairs_per_unit=repairs_per_unit, common_utilities=common_utilities,
                leasing_reserve=leasing_reserve,
                exit_cost_pct=exit_cost_pct, price_decline=price_decline,
                split_soft=split_soft, arch_pct=arch_pct, eng_pct=eng_pct,
                permit_fee_pct=permit_fee_pct, survey_pct=survey_pct,
                insurance_dev_pct=insurance_dev_pct, other_soft_pct=other_soft_pct,
                broker_fee_pct=broker_fee_pct, title_closing_pct=title_closing_pct,
                seller_concessions_pct=seller_concessions_pct,
                hard_cost=hard_cost, hard_contingency=hard_contingency, soft_costs=soft_costs,
                total_dev_cost=total_dev_cost, total_project_cost=total_project_cost,
                loan_amount=loan_amount, equity=equity,
                construction_interest=construction_interest, loan_fees=loan_fees,
                hold_interest=hold_interest, gross_rent=gross_rent,
                effective_rent=effective_rent,
                total_hold_expenses=total_hold_expenses, net_rental_income=net_rental_income,
                monthly_debt_service=monthly_debt_service,
                loan_balance_after_hold=loan_balance_after_hold,
                total_equity_invested=total_equity_invested,
                user_profit=user_profit, user_revenue=user_revenue,
                breakeven_psf=breakeven_psf, exit_costs_user=exit_costs_user,
                total_cost=total_cost, equity_multiple=equity_multiple,
                additional_equity_needed=additional_equity_needed,
                cumulative_cf=cumulative_cf,
                net_sale_before_debt=_nsbd,
                net_sale_after_debt=_nsad,
                positive_rental_cf=_prcf,
                total_cash_returned=total_cash_returned,
                mgmt_cost=mgmt_cost, prop_tax=prop_tax, insurance=insurance,
                repairs=repairs, misc=misc, leasing=leasing,
                monthly_noi=monthly_noi, monthly_cf_after_debt=monthly_cf_after_debt,
            )
                st.download_button(
                    label="📥 Download Excel Model (.xlsx)",
                    data=excel_bytes,
                    file_name=f"{street_name.title()}_Hold_Model.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                    key="dl_excel",
                )
                st.caption("Contains 2 sheets: **Hold Model** + **Detailed Analysis**")
            except Exception as e:
                st.error(f"Excel generation error: {e}")
        else:
            st.warning("No data to generate report.")

        st.markdown("---")
        st.subheader("📊 Key Insights")
        st.markdown(f"""
        - **Break-even exit:** ${breakeven_psf:.0f}/sf
        - **Market median:** ${median_psf}/sf (Redfin, {result.market_stats.get('count', 0)} comps)
        - **Your exit:** ${exit_psf}/sf {'✅' if exit_psf <= median_psf else '⚠️'}
        - **This deal works IF:**
          - Purchase price ≤ ${max(0, purchase_price - (exit_psf - median_psf) * build_sf):,.0f} (or lower)
          - Exit ≥ ${breakeven_psf:.0f}/sf
          - No major delays beyond {delay_months} months
        """)

elif submitted and not show_analysis:
    # No address/ZIP — run financial-only mode
    st.info("💡 No address entered — showing financial analysis only. Add address + ZIP for market data, comps, and permits.")

    # Create a dummy result with no market data
    class EmptyResult:
        street_permits = []
        zip_permits = []
        redfin_comps = []
        market_stats = {}
        sources_status = {}
    result = EmptyResult()

    # ── Financial calculations (same as full mode) ──
    hard_cost = build_cost_psf * build_sf
    hard_contingency = hard_cost * (hard_contingency_pct / 100)
    soft_costs = hard_cost * (soft_cost_pct / 100)
    soft_contingency = 0
    total_dev_cost = hard_cost + hard_contingency + soft_costs + soft_contingency
    total_project_cost = purchase_price + total_dev_cost

    loan_amount = total_dev_cost * (ltv / 100)
    equity = total_project_cost - loan_amount
    total_build_months = build_months + delay_months
    construction_interest = loan_amount * (interest_rate / 100) * (total_build_months / 12) * (draw_factor / 100)
    loan_fees = loan_amount * (loan_fee_pct / 100)

    hold_interest = 0
    gross_rent = 0
    effective_rent = 0
    total_hold_expenses = 0
    net_rental_income = 0
    monthly_debt_service = 0
    loan_balance_after_hold = loan_amount
    if hold_months > 0:
        monthly_rent_total = rent_per_unit * units
        gross_rent = monthly_rent_total * hold_months
        effective_rent = gross_rent * (1 - vacancy_pct / 100)
        mgmt_expense = effective_rent * (mgmt_fee_pct / 100)
        prop_tax = build_sf * taxable_value_psf * (prop_tax_rate / 100) * hold_months / 12
        insurance = insurance_monthly * hold_months
        repairs = repairs_per_unit * units * hold_months
        misc = common_utilities * hold_months
        leasing = leasing_reserve * hold_months
        # Permanent mortgage debt service
        monthly_perm = perm_mortgage_rate / 100 / 12
        n_pay = amortization_years * 12
        if monthly_perm > 0:
            monthly_debt_service = loan_amount * (monthly_perm * (1 + monthly_perm) ** n_pay) / ((1 + monthly_perm) ** n_pay - 1)
            loan_balance_after_hold = loan_amount * (1 + monthly_perm) ** hold_months - monthly_debt_service * ((1 + monthly_perm) ** hold_months - 1) / monthly_perm
        else:
            monthly_debt_service = loan_amount / n_pay
            loan_balance_after_hold = loan_amount - (monthly_debt_service * hold_months)
        hold_interest = monthly_debt_service * hold_months
        total_hold_expenses = mgmt_expense + prop_tax + insurance + repairs + misc + leasing + hold_interest
        net_rental_income = effective_rent - total_hold_expenses

    total_interest = construction_interest + hold_interest
    timeline_years = (build_months + hold_months + delay_months) / 12

    median_psf = 0
    market_exit = exit_psf

    price_change_rate = price_decline / 100
    adjusted_exit = market_exit * ((1 + price_change_rate) ** timeline_years)
    adjusted_revenue = adjusted_exit * build_sf
    user_revenue = exit_psf * build_sf

    exit_costs_user = user_revenue * (exit_cost_pct / 100)
    exit_costs_market = adjusted_revenue * (exit_cost_pct / 100)
    total_cost = total_project_cost + total_interest + exit_costs_user
    market_total_cost = total_project_cost + total_interest + exit_costs_market
    market_profit = adjusted_revenue - market_total_cost + net_rental_income
    user_profit = user_revenue - total_cost + net_rental_income

    annual_rent = rent_per_unit * 12 * units
    cash_yield = (annual_rent / equity * 100) if equity > 0 else 0
    annualized_return = ((user_revenue / total_cost) ** (1 / max(timeline_years, 0.5)) - 1) if total_cost > 0 else 0
    breakeven_psf = (total_cost - net_rental_income) / build_sf if build_sf > 0 else 0

    # ── Display financial results ──
    st.subheader("📊 Deal Snapshot")
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Total All-In Cost", f"${total_cost:,.0f}")
    col2.metric("Sale Revenue", f"${user_revenue:,.0f}")
    col3.metric("Profit", f"${user_profit:,.0f}", delta=f"{annualized_return*100:.1f}% ann.")
    col4.metric("Break-Even", f"${breakeven_psf:.0f}/sf")

    if user_profit > 200000:
        st.success(f"✅ STRONG — ${user_profit:,.0f} profit at ${exit_psf}/sf exit")
    elif user_profit > 50000:
        st.warning(f"⚠️ MARGINAL — ${user_profit:,.0f} profit, sensitive to delays/overruns")
    else:
        st.error(f"❌ WEAK/LOSS — ${user_profit:,.0f} at ${exit_psf}/sf exit")

    # Cost breakdown
    st.subheader("💰 Cost Breakdown")
    soft_line_fin = ""
    if split_soft:
        soft_line_fin = f"""| — Architecture & Design ({arch_pct}%) | ${hard_cost * arch_pct / 100:,.0f} |
    | — Engineering ({eng_pct}%) | ${hard_cost * eng_pct / 100:,.0f} |
    | — Permits & Impact Fees ({permit_fee_pct}%) | ${hard_cost * permit_fee_pct / 100:,.0f} |
    | — Surveys & Geotech ({survey_pct}%) | ${hard_cost * survey_pct / 100:,.0f} |
    | — Builder's Risk Insurance ({insurance_dev_pct}%) | ${hard_cost * insurance_dev_pct / 100:,.0f} |
    | — Other Soft ({other_soft_pct}%) | ${hard_cost * other_soft_pct / 100:,.0f} |
    | **Soft Costs Total ({soft_cost_pct:.1f}%)** | **${soft_costs:,.0f}** |"""
    else:
        soft_line_fin = f"| Soft Costs ({soft_cost_pct}%) | ${soft_costs:,.0f} |"
    st.markdown(f"""
    | Item | Amount |
    |------|--------|
    | Land / Purchase | ${purchase_price:,.0f} |
    | Hard Cost ({build_sf:,} sf × ${build_cost_psf}/sf) | ${hard_cost:,.0f} |
    | Hard Contingency ({hard_contingency_pct}%) | ${hard_contingency:,.0f} |
    {soft_line_fin}
    | Construction Interest | ${construction_interest:,.0f} |
    | Hold Debt Service ({hold_months} mo) | ${hold_interest:,.0f} |
    | **Sale Costs ({exit_cost_pct:.1f}%)** | **${exit_costs_user:,.0f}** |
    | — Broker/Agent ({broker_fee_pct}%) | ${user_revenue * broker_fee_pct / 100:,.0f} |
    | — Title + Closing ({title_closing_pct}%) | ${user_revenue * title_closing_pct / 100:,.0f} |
    | — Seller Concessions ({seller_concessions_pct}%) | ${user_revenue * seller_concessions_pct / 100:,.0f} |
    | **Total All-In** | **${total_cost:,.0f}** |
    """)

    if hold_months > 0:
        st.subheader("🏠 Rental Hold Detail")
        st.markdown(f"""
        | Item | Monthly | Total ({hold_months} mo) |
        |------|---------|---------|
        | Gross Rent ({units} × ${rent_per_unit:,}/mo) | ${rent_per_unit * units:,.0f} | ${gross_rent:,.0f} |
        | Less Vacancy ({vacancy_pct}%) | -${rent_per_unit * units * vacancy_pct / 100:,.0f} | -${gross_rent * vacancy_pct / 100:,.0f} |
        | Less Mgmt ({mgmt_fee_pct}%) | | -${effective_rent * mgmt_fee_pct / 100:,.0f} |
        | Less Property Tax ({prop_tax_rate}%) | | -${prop_tax:,.0f} |
        | Less Insurance/Repairs/Misc/Leasing | | -${insurance + repairs + misc + leasing:,.0f} |
        | Less Debt Service ({perm_mortgage_rate}%, {amortization_years}yr) | -${monthly_debt_service:,.0f} | -${hold_interest:,.0f} |
        | **Net Rental Income** | | **${net_rental_income:,.0f}** |
        """)

    # Sensitivity
    st.subheader("📈 Sensitivity — Exit $/sf vs Profit")
    exit_range = np.arange(max(200, int(breakeven_psf) - 100), int(breakeven_psf) + 200, 10)
    profits_by_exit = []
    for e in exit_range:
        rev = e * build_sf
        exit_c = rev * (exit_cost_pct / 100)
        profits_by_exit.append(rev - (total_project_cost + total_interest + exit_c) + net_rental_income)
    st.line_chart({"Exit $/sf": exit_range, "Profit": profits_by_exit}, x="Exit $/sf", y="Profit")
elif not show_analysis:
    # Landing page
    st.markdown("""
    ### 🏡 How It Works

    1. **Enter property details** in the sidebar (address, ZIP, deal numbers)
    2. **Click "Analyze"** — the app pulls real market data automatically
    3. **Get a clear BUY / CAUTION / DON'T BUY verdict**

    ### What Makes This Different

    | Feature | Generic Calculator | This App |
    |---------|-------------------|----------|
    | Market comps | ❌ You guess | ✅ Pulls from Redfin |
    | Permit data | ❌ None | ✅ Austin permits API |
    | Competition check | ❌ None | ✅ Same-street construction |
    | Risk scoring | ❌ None | ✅ Automated 0-100 score |
    | BUY/DON'T BUY | ❌ None | ✅ Clear recommendation |
    | Sensitivity analysis | ✅ Manual | ✅ Auto with real median |

    > **Pro tip:** For complete analysis including TCAD records, deed history,
    > and foreclosure detection, run the CLI tool locally:
    > ```
    > python analyze_deal.py "ADDRESS" --zip XXXXX --purchase-price N
    > ```
    """)
